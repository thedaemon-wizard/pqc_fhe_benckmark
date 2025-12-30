#!/usr/bin/env python3
"""
Generate Complete Word Document for PQC-FHE Technical Report v2.3.4
Matches PDF version exactly with all sections
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_table_with_header(doc, headers, data, header_color):
    """Create a table with colored header row"""
    table = doc.add_table(rows=len(data)+1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Header row
    header_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        header_cells[i].text = h
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(header_cells[i], header_color)
    
    # Data rows
    for i, row_data in enumerate(data):
        row = table.rows[i + 1]
        for j, val in enumerate(row_data):
            row.cells[j].text = str(val)
    
    return table

def create_report():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    # Configure Heading styles
    h1_style = doc.styles['Heading 1']
    h1_style.font.name = 'Arial'
    h1_style.font.size = Pt(16)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0x2c, 0x52, 0x82)
    
    h2_style = doc.styles['Heading 2']
    h2_style.font.name = 'Arial'
    h2_style.font.size = Pt(12)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(0x31, 0x82, 0xce)
    
    # =========================================================================
    # TITLE PAGE
    # =========================================================================
    doc.add_paragraph()
    doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PQC-FHE Integration Platform")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Technical Report v2.3.4")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x4a, 0x55, 0x68)
    
    subtitle2 = doc.add_paragraph()
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle2.add_run("Post-Quantum Cryptography + Fully Homomorphic Encryption")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x4a, 0x55, 0x68)
    
    doc.add_paragraph()
    
    # Info table
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = 'Table Grid'
    info_data = [
        ('Version', '2.3.4'),
        ('Release Date', '2025-12-30'),
        ('PQC Standards', 'FIPS 203 (ML-KEM), FIPS 204 (ML-DSA)'),
        ('FHE Scheme', 'CKKS (DESILO Implementation)'),
        ('License', 'MIT'),
    ]
    for i, (label, value) in enumerate(info_data):
        row = info_table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x2c, 0x52, 0x82)
        row.cells[1].text = value
    
    doc.add_page_break()
    
    # =========================================================================
    # TABLE OF CONTENTS
    # =========================================================================
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        "1. Executive Summary",
        "2. Market Context and Business Value",
        "3. System Architecture",
        "4. Post-Quantum Cryptography Implementation",
        "5. Fully Homomorphic Encryption Implementation",
        "6. Live Data Sources",
        "7. Enterprise Use Cases",
        "8. API Reference",
        "9. Installation Guide (Multi-Platform)",
        "10. Security Analysis",
        "11. Future Roadmap",
        "Appendix A: Algorithm Parameters",
        "Appendix B: Performance Benchmarks",
        "References",
    ]
    for item in toc_items:
        doc.add_paragraph(item)
    
    doc.add_page_break()
    
    # =========================================================================
    # 1. EXECUTIVE SUMMARY
    # =========================================================================
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        "The PQC-FHE Integration Platform provides a production-ready framework combining "
        "Post-Quantum Cryptography (PQC) with Fully Homomorphic Encryption (FHE) for "
        "enterprise security applications. This platform addresses the emerging threat of "
        "quantum computers to current cryptographic systems while enabling privacy-preserving "
        "computation on sensitive data."
    )
    
    doc.add_heading('Key Features', level=2)
    features = [
        "NIST-standardized PQC algorithms (FIPS 203, FIPS 204)",
        "CKKS-based homomorphic encryption via DESILO FHE",
        "Real-time data integration from verified public sources",
        "REST API with comprehensive Swagger documentation",
        "Interactive Web UI with live demonstrations",
        "GPU acceleration support (CUDA 12.x/13.x)",
    ]
    for f in features:
        doc.add_paragraph(f, style='List Bullet')
    
    doc.add_heading('v2.3.4 Updates (2025-12-30)', level=2)
    doc.add_paragraph(
        "Version 2.3.4 introduces enhanced live data fetching capabilities with robust "
        "fallback mechanisms, fixes for numpy array handling in FHE demo endpoints, "
        "and improved multi-platform installation documentation."
    )
    
    doc.add_page_break()
    
    # =========================================================================
    # 2. MARKET CONTEXT
    # =========================================================================
    doc.add_heading('2. Market Context and Business Value', level=1)
    
    doc.add_heading('The Quantum Threat', level=2)
    doc.add_paragraph(
        "Quantum computers pose an existential threat to current public-key cryptography. "
        "Shor's algorithm can break RSA-2048 in polynomial time, and Grover's algorithm "
        "reduces symmetric key security by half. NIST estimates cryptographically-relevant "
        "quantum computers may emerge within 10-15 years, necessitating immediate migration "
        "to quantum-resistant alternatives."
    )
    
    doc.add_heading('Market Opportunity', level=2)
    add_table_with_header(doc,
        ['Segment', '2024', '2034', 'CAGR'],
        [
            ('PQC Solutions', '$302M', '$30B', '58%'),
            ('FHE Market', '$200M', '$2.5B', '28%'),
            ('Quantum Security', '$1.2B', '$15B', '29%'),
        ],
        '2c5282')
    
    p = doc.add_paragraph()
    run = p.add_run("Source: Markets and Markets, Gartner (2024)")
    run.font.italic = True
    run.font.size = Pt(9)
    
    doc.add_page_break()
    
    # =========================================================================
    # 3. SYSTEM ARCHITECTURE
    # =========================================================================
    doc.add_heading('3. System Architecture', level=1)
    
    doc.add_heading('Component Overview', level=2)
    add_table_with_header(doc,
        ['Layer', 'Component', 'Technology'],
        [
            ('Presentation', 'Web UI', 'React + Tailwind CSS'),
            ('API', 'REST Server', 'FastAPI + Swagger UI'),
            ('Cryptography', 'PQC Manager', 'liboqs-python (ML-KEM, ML-DSA)'),
            ('Cryptography', 'FHE Engine', 'DESILO FHE (CKKS)'),
            ('Data', 'Live Fetcher', 'VitalDB, yfinance, Ethereum RPC'),
            ('Infrastructure', 'GPU Accel.', 'CUDA 12.x/13.x + cuQuantum'),
        ],
        '3182ce')
    
    doc.add_heading('Data Flow', level=2)
    flow_steps = [
        "Client sends request to REST API (FastAPI server on port 8000)",
        "API validates input and routes to appropriate handler",
        "For PQC operations: liboqs-python performs key generation/signing",
        "For FHE operations: DESILO engine encrypts/computes/decrypts",
        "Live data fetcher retrieves external data with automatic fallback",
        "Response returned to client with full audit trail",
    ]
    for i, step in enumerate(flow_steps, 1):
        doc.add_paragraph(f"{i}. {step}")
    
    doc.add_page_break()
    
    # =========================================================================
    # 4. PQC IMPLEMENTATION
    # =========================================================================
    doc.add_heading('4. Post-Quantum Cryptography Implementation', level=1)
    
    doc.add_heading('NIST Standards Compliance', level=2)
    doc.add_paragraph(
        "This platform implements NIST's finalized post-quantum cryptography standards: "
        "FIPS 203 (ML-KEM) for key encapsulation and FIPS 204 (ML-DSA) for digital "
        "signatures. These standards were published on August 13, 2024 and represent "
        "the foundation of quantum-resistant cryptography."
    )
    
    doc.add_heading('Key Encapsulation Mechanisms (FIPS 203)', level=2)
    add_table_with_header(doc,
        ['Algorithm', 'Level', 'Public Key', 'Ciphertext', 'Use Case'],
        [
            ('ML-KEM-512', '1', '800 B', '768 B', 'IoT/Embedded'),
            ('ML-KEM-768', '3', '1,184 B', '1,088 B', 'General Purpose'),
            ('ML-KEM-1024', '5', '1,568 B', '1,568 B', 'High Security'),
        ],
        '48bb78')
    
    doc.add_paragraph()
    
    doc.add_heading('Digital Signature Algorithms (FIPS 204)', level=2)
    add_table_with_header(doc,
        ['Algorithm', 'Level', 'Public Key', 'Signature', 'Speed'],
        [
            ('ML-DSA-44', '2', '1,312 B', '2,420 B', 'Fastest'),
            ('ML-DSA-65', '3', '1,952 B', '3,309 B', 'Balanced'),
            ('ML-DSA-87', '5', '2,592 B', '4,627 B', 'Maximum'),
        ],
        'ed8936')
    
    doc.add_page_break()
    
    # =========================================================================
    # 5. FHE IMPLEMENTATION
    # =========================================================================
    doc.add_heading('5. Fully Homomorphic Encryption Implementation', level=1)
    
    doc.add_heading('CKKS Scheme Overview', level=2)
    doc.add_paragraph(
        "The platform uses the CKKS (Cheon-Kim-Kim-Song) homomorphic encryption scheme, "
        "which supports approximate arithmetic on encrypted real numbers. CKKS is ideal "
        "for machine learning and statistical analysis on encrypted data [4]."
    )
    
    doc.add_heading('Key Properties', level=2)
    props = [
        "Supports addition and multiplication on encrypted floating-point numbers",
        "Slot packing allows SIMD operations on thousands of values simultaneously",
        "Rescaling maintains precision across multiplicative depth",
        "Bootstrapping enables unlimited computation depth",
    ]
    for p in props:
        doc.add_paragraph(p, style='List Bullet')
    
    doc.add_heading('DESILO FHE Configuration', level=2)
    add_table_with_header(doc,
        ['Parameter', 'Value', 'Description'],
        [
            ('poly_degree', '16,384', 'Polynomial ring dimension (N)'),
            ('coeff_mod_bit_sizes', '[60,40,40,40,60]', 'Coefficient modulus chain'),
            ('scale', '2^40', 'Encoding scale for precision'),
            ('max_mult_depth', '4', 'Maximum multiplicative depth'),
            ('slot_count', '8,192', 'Number of plaintext slots'),
        ],
        '805ad5')
    
    doc.add_page_break()
    
    # =========================================================================
    # 6. LIVE DATA SOURCES
    # =========================================================================
    doc.add_heading('6. Live Data Sources', level=1)
    doc.add_paragraph(
        "Version 2.3.4 provides robust live data fetching with automatic fallback to "
        "embedded sample data. This ensures demonstrations work reliably while "
        "showcasing real-world data integration capabilities."
    )
    
    doc.add_heading('Healthcare: VitalDB [6]', level=2)
    add_table_with_header(doc,
        ['Property', 'Value'],
        [
            ('Dataset', 'VitalDB Open Dataset'),
            ('Method', 'vitaldb Python library'),
            ('Data Type', 'Surgical patient vital signs (BP, HR, SpO2)'),
            ('Sample Size', '6,388 surgical cases'),
            ('DOI', '10.1038/s41597-022-01411-5'),
            ('License', 'CC BY-NC-SA 4.0'),
        ],
        'e53e3e')
    
    doc.add_paragraph()
    
    doc.add_heading('Finance: Yahoo Finance [8]', level=2)
    add_table_with_header(doc,
        ['Property', 'Value'],
        [
            ('Method', 'yfinance Python library'),
            ('Data Type', 'Real-time stock prices, market cap'),
            ('Symbols', 'AAPL, MSFT, GOOGL, AMZN, NVDA, META, TSLA, JPM'),
            ('License', 'Yahoo Finance Terms of Service'),
        ],
        '38a169')
    
    doc.add_paragraph()
    
    doc.add_heading('IoT: UCI Machine Learning Repository [7]', level=2)
    add_table_with_header(doc,
        ['Property', 'Value'],
        [
            ('Dataset', 'Individual Household Electric Power Consumption'),
            ('Data Type', 'Smart meter power readings'),
            ('DOI', '10.24432/C52G6F'),
            ('License', 'CC BY 4.0'),
        ],
        '3182ce')
    
    doc.add_paragraph()
    
    doc.add_heading('Blockchain: Ethereum RPC Endpoints [9]', level=2)
    add_table_with_header(doc,
        ['Priority', 'Endpoint', 'Provider'],
        [
            ('1', 'rpc.ankr.com/eth', 'Ankr (Primary)'),
            ('2', 'ethereum-rpc.publicnode.com', 'PublicNode'),
            ('3', 'cloudflare-eth.com', 'Cloudflare'),
            ('4', 'eth.drpc.org', 'DRPC'),
            ('5', '1rpc.io/eth', '1RPC'),
        ],
        '4a5568')
    
    p = doc.add_paragraph()
    run = p.add_run("Note: All endpoints are free and require no API key.")
    run.font.italic = True
    
    doc.add_page_break()
    
    # =========================================================================
    # 7. ENTERPRISE USE CASES
    # =========================================================================
    doc.add_heading('7. Enterprise Use Cases', level=1)
    
    doc.add_heading('Healthcare: HIPAA-Compliant Analytics', level=2)
    doc.add_paragraph(
        "Hospitals can analyze patient vital signs without exposing Protected Health "
        "Information (PHI). FHE enables computation on encrypted blood pressure readings "
        "to identify hypertension trends across populations while maintaining full HIPAA "
        "compliance. Clinical interpretation follows AHA Guidelines [12]."
    )
    
    doc.add_heading('Finance: Confidential Portfolio Analysis', level=2)
    doc.add_paragraph(
        "Investment firms can perform growth projections on encrypted portfolio values. "
        "Client holdings remain confidential even during third-party analysis, enabling "
        "secure outsourcing of financial computations."
    )
    
    doc.add_heading('IoT: Secure Smart Grid Analytics', level=2)
    doc.add_paragraph(
        "Utility companies can aggregate encrypted smart meter readings for demand "
        "forecasting without accessing individual household consumption patterns, "
        "preserving consumer privacy while enabling grid optimization."
    )
    
    doc.add_heading('Blockchain: Quantum-Resistant Transactions', level=2)
    doc.add_paragraph(
        "Cryptocurrency platforms can migrate from ECDSA to ML-DSA signatures, protecting "
        "transaction integrity against future quantum attacks. This platform demonstrates "
        "the migration path with side-by-side comparison of signature sizes."
    )
    
    doc.add_page_break()
    
    # =========================================================================
    # 8. API REFERENCE
    # =========================================================================
    doc.add_heading('8. API Reference', level=1)
    
    doc.add_heading('Endpoint Summary', level=2)
    add_table_with_header(doc,
        ['Endpoint', 'Method', 'Description'],
        [
            ('/health', 'GET', 'Health check'),
            ('/pqc/algorithms', 'GET', 'List PQC algorithms'),
            ('/pqc/kem/keypair', 'POST', 'Generate KEM keypair'),
            ('/pqc/kem/encapsulate', 'POST', 'Encapsulate secret'),
            ('/pqc/kem/decapsulate', 'POST', 'Decapsulate secret'),
            ('/pqc/sig/keypair', 'POST', 'Generate SIG keypair'),
            ('/pqc/sig/sign', 'POST', 'Sign message'),
            ('/pqc/sig/verify', 'POST', 'Verify signature'),
            ('/fhe/encrypt', 'POST', 'Encrypt data'),
            ('/fhe/decrypt', 'POST', 'Decrypt ciphertext'),
            ('/fhe/add', 'POST', 'Add ciphertexts'),
            ('/fhe/multiply', 'POST', 'Multiply by scalar'),
            ('/enterprise/healthcare', 'GET', 'Healthcare data'),
            ('/enterprise/finance', 'GET', 'Finance data'),
            ('/enterprise/iot', 'GET', 'IoT data'),
            ('/enterprise/blockchain', 'GET', 'Blockchain data'),
        ],
        '2c5282')
    
    doc.add_paragraph()
    doc.add_paragraph("Interactive documentation available at http://localhost:8000/docs (Swagger UI)")
    
    doc.add_page_break()
    
    # =========================================================================
    # 9. INSTALLATION GUIDE
    # =========================================================================
    doc.add_heading('9. Installation Guide (Multi-Platform)', level=1)
    
    doc.add_heading('System Requirements', level=2)
    add_table_with_header(doc,
        ['Component', 'Minimum', 'Recommended'],
        [
            ('Python', '3.9+', '3.11+'),
            ('RAM', '8 GB', '32 GB'),
            ('Storage', '1 GB', '5 GB'),
            ('GPU (optional)', 'CUDA 11.x', 'CUDA 12.x/13.x'),
        ],
        '4a5568')
    
    doc.add_heading('Build Dependencies by Platform', level=2)
    
    platforms = [
        ("Debian/Ubuntu:", "sudo apt update\nsudo apt install -y cmake gcc g++ libssl-dev python3-dev git"),
        ("Fedora/RHEL/CentOS:", "sudo dnf install -y cmake gcc gcc-c++ openssl-devel python3-devel git"),
        ("Arch Linux:", "sudo pacman -S cmake gcc openssl python git"),
        ("macOS (Homebrew):", "brew install cmake openssl@3 git"),
    ]
    
    for platform_name, cmd in platforms:
        p = doc.add_paragraph()
        run = p.add_run(platform_name)
        run.font.bold = True
        
        p = doc.add_paragraph()
        run = p.add_run(cmd)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        doc.add_paragraph()
    
    doc.add_heading('Installing liboqs-python [10]', level=2)
    doc.add_paragraph(
        "liboqs-python is NOT available via pip. It must be built from source. "
        "The recommended method auto-downloads and builds liboqs at runtime:"
    )
    
    p = doc.add_paragraph()
    run = p.add_run("Option A: Automatic Build (Recommended)")
    run.font.bold = True
    
    code_a = """git clone --depth=1 https://github.com/open-quantum-safe/liboqs-python
cd liboqs-python
pip install .
cd .."""
    p = doc.add_paragraph()
    run = p.add_run(code_a)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Option B: Manual Build")
    run.font.bold = True
    
    code_b = """# Build liboqs C library
git clone --depth=1 https://github.com/open-quantum-safe/liboqs
cmake -S liboqs -B liboqs/build -DBUILD_SHARED_LIBS=ON
cmake --build liboqs/build --parallel 8
sudo cmake --build liboqs/build --target install

# Set library path (Linux)
export LD_LIBRARY_PATH=$LD_LIBRARY_PATH:/usr/local/lib

# Install Python wrapper
git clone --depth=1 https://github.com/open-quantum-safe/liboqs-python
cd liboqs-python && pip install . && cd .."""
    p = doc.add_paragraph()
    run = p.add_run(code_b)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    
    doc.add_page_break()
    
    doc.add_heading('Installing DESILO FHE [5]', level=2)
    code_fhe = """# CPU mode
pip install desilofhe

# GPU mode (choose based on CUDA version)
pip install desilofhe-cu130  # CUDA 13.0
pip install desilofhe-cu124  # CUDA 12.4
pip install desilofhe-cu121  # CUDA 12.1"""
    p = doc.add_paragraph()
    run = p.add_run(code_fhe)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    
    doc.add_heading('Complete Quick Start', level=2)
    code_quick = """# 1. Clone repository
git clone https://github.com/thedaemon-wizard/pqc_fhe_benckmark.git
cd pqc_fhe_benckmark
pip install -r requirements.txt

# 2. Install liboqs-python
git clone --depth=1 https://github.com/open-quantum-safe/liboqs-python
cd liboqs-python && pip install . && cd ..

# 3. Install DESILO FHE
pip install desilofhe

# 4. Optional: Live data libraries
pip install yfinance vitaldb

# 5. Start server
python -m uvicorn api.server:app --host 0.0.0.0 --port 8000

# 6. Access Web UI
# Open http://localhost:8000/ui"""
    p = doc.add_paragraph()
    run = p.add_run(code_quick)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    
    doc.add_page_break()
    
    # =========================================================================
    # 10. SECURITY ANALYSIS
    # =========================================================================
    doc.add_heading('10. Security Analysis', level=1)
    
    doc.add_heading('NIST Security Levels', level=2)
    doc.add_paragraph(
        "NIST defines five security levels based on classical and quantum attack "
        "complexity. This platform supports Levels 1, 2, 3, and 5:"
    )
    
    add_table_with_header(doc,
        ['Level', 'Classical Equivalent', 'Algorithms'],
        [
            ('1', 'AES-128', 'ML-KEM-512'),
            ('2', 'SHA-256', 'ML-DSA-44'),
            ('3', 'AES-192', 'ML-KEM-768, ML-DSA-65'),
            ('5', 'AES-256', 'ML-KEM-1024, ML-DSA-87'),
        ],
        'e53e3e')
    
    doc.add_heading('FHE Security', level=2)
    doc.add_paragraph(
        "CKKS security is based on the Ring Learning With Errors (RLWE) problem, which "
        "is believed to be hard for both classical and quantum computers. The configured "
        "parameters provide at least 128-bit security against known attacks."
    )
    
    doc.add_page_break()
    
    # =========================================================================
    # 11. FUTURE ROADMAP
    # =========================================================================
    doc.add_heading('11. Future Roadmap', level=1)
    
    add_table_with_header(doc,
        ['Version', 'Timeline', 'Features'],
        [
            ('v2.4.0', 'Q1 2025', 'SLH-DSA (SPHINCS+) hash-based signatures'),
            ('v2.5.0', 'Q2 2025', 'Hybrid classical/PQC mode'),
            ('v2.6.0', 'Q3 2025', 'Multi-party computation (MPC) integration'),
            ('v3.0.0', 'Q4 2025', 'FIPS validation and enterprise hardening'),
        ],
        '38a169')
    
    doc.add_page_break()
    
    # =========================================================================
    # APPENDIX A: ALGORITHM PARAMETERS
    # =========================================================================
    doc.add_heading('Appendix A: Algorithm Parameters', level=1)
    
    doc.add_heading('ML-KEM Parameters (FIPS 203)', level=2)
    add_table_with_header(doc,
        ['Parameter', 'ML-KEM-512', 'ML-KEM-768', 'ML-KEM-1024'],
        [
            ('n (dimension)', '256', '256', '256'),
            ('k', '2', '3', '4'),
            ('eta_1', '3', '2', '2'),
            ('eta_2', '2', '2', '2'),
            ('d_u', '10', '10', '11'),
            ('d_v', '4', '4', '5'),
        ],
        '4299e1')
    
    doc.add_paragraph()
    
    doc.add_heading('ML-DSA Parameters (FIPS 204)', level=2)
    add_table_with_header(doc,
        ['Parameter', 'ML-DSA-44', 'ML-DSA-65', 'ML-DSA-87'],
        [
            ('q (modulus)', '8,380,417', '8,380,417', '8,380,417'),
            ('k', '4', '6', '8'),
            ('l', '4', '5', '7'),
            ('eta', '2', '4', '2'),
            ('tau', '39', '49', '60'),
        ],
        'ed8936')
    
    doc.add_page_break()
    
    # =========================================================================
    # APPENDIX B: PERFORMANCE BENCHMARKS
    # =========================================================================
    doc.add_heading('Appendix B: Performance Benchmarks', level=1)
    
    doc.add_heading('PQC Operations (Intel i7-12700H)', level=2)
    add_table_with_header(doc,
        ['Operation', 'ML-KEM-768', 'ML-DSA-65'],
        [
            ('Key Generation', '0.03 ms', '0.08 ms'),
            ('Encap/Sign', '0.04 ms', '0.18 ms'),
            ('Decap/Verify', '0.04 ms', '0.06 ms'),
        ],
        '667eea')
    
    doc.add_paragraph()
    
    doc.add_heading('FHE Operations', level=2)
    add_table_with_header(doc,
        ['Operation', 'CPU', 'GPU (RTX 4090)'],
        [
            ('Key Generation', '2.5 s', '0.8 s'),
            ('Encrypt (8192 slots)', '15 ms', '3 ms'),
            ('Add', '0.5 ms', '0.1 ms'),
            ('Multiply (scalar)', '2 ms', '0.3 ms'),
            ('Multiply (ct*ct)', '50 ms', '8 ms'),
            ('Decrypt', '10 ms', '2 ms'),
        ],
        '9f7aea')
    
    doc.add_page_break()
    
    # =========================================================================
    # REFERENCES
    # =========================================================================
    doc.add_heading('References', level=1)
    
    references = [
        "[1] NIST. FIPS 203: Module-Lattice-Based Key-Encapsulation Mechanism Standard. National Institute of Standards and Technology, August 2024. https://csrc.nist.gov/pubs/fips/203/final",
        "[2] NIST. FIPS 204: Module-Lattice-Based Digital Signature Standard. National Institute of Standards and Technology, August 2024. https://csrc.nist.gov/pubs/fips/204/final",
        "[3] NIST. FIPS 205: Stateless Hash-Based Digital Signature Standard. National Institute of Standards and Technology, August 2024. https://csrc.nist.gov/pubs/fips/205/final",
        "[4] Cheon JH, Kim A, Kim M, Song Y. Homomorphic Encryption for Arithmetic of Approximate Numbers. ASIACRYPT 2017. DOI: 10.1007/978-3-319-70694-8_15",
        "[5] DESILO. DESILO FHE Library Documentation. https://fhe.desilo.dev/latest/",
        "[6] Lee HC, Park Y, Yoon SB, et al. VitalDB, a high-fidelity multi-parameter vital signs database in surgical patients. Scientific Data 9, 279 (2022). DOI: 10.1038/s41597-022-01411-5",
        "[7] Hebrail G, Berard A. Individual Household Electric Power Consumption Data Set. UCI Machine Learning Repository. DOI: 10.24432/C52G6F",
        "[8] Aroussi R. yfinance: Download market data from Yahoo! Finance API. https://github.com/ranaroussi/yfinance. License: Apache 2.0",
        "[9] Ethereum Foundation. Ethereum JSON-RPC API. https://ethereum.org/developers/docs/apis/json-rpc/",
        "[10] Open Quantum Safe Project. liboqs-python: Python 3 wrapper for liboqs. https://github.com/open-quantum-safe/liboqs-python. License: MIT",
        "[11] Ramirez S, et al. FastAPI: Modern, fast web framework for building APIs. https://fastapi.tiangolo.com/. License: MIT",
        "[12] American Heart Association. Understanding Blood Pressure Readings. https://www.heart.org/en/health-topics/high-blood-pressure/understanding-blood-pressure-readings",
    ]
    
    for ref in references:
        p = doc.add_paragraph(ref)
        p.paragraph_format.space_after = Pt(6)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("---")
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Generated: 2025-12-30 | PQC-FHE Integration Platform v2.3.4")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    # Save document
    doc.save("PQC_FHE_Technical_Report_v2.3.4.docx")
    print("Word document generated: PQC_FHE_Technical_Report_v2.3.4.docx")

if __name__ == "__main__":
    create_report()
