#!/usr/bin/env python3
"""
Update PQC-FHE Technical Report from v2.3.5 to v3.0.0
Adds new sections: Quantum Threat Simulator, Security Scoring, MPC-HE, Extended Benchmarks
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os


def set_cell_shading(cell, color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_styled_table(doc, headers, data, header_color, widths=None):
    """Create a professionally styled table"""
    table = doc.add_table(rows=len(data)+1, cols=len(headers))
    table.style = 'Table Grid'

    if widths:
        for i, width in enumerate(widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)

    header_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        header_cells[i].text = h
        for paragraph in header_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(header_cells[i], header_color)

    for i, row_data in enumerate(data):
        row = table.rows[i + 1]
        for j, val in enumerate(row_data):
            row.cells[j].text = str(val)
            for paragraph in row.cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
        if i % 2 == 1:
            for cell in row.cells:
                set_cell_shading(cell, 'F7FAFC')

    return table


def add_heading(doc, text, level=1):
    """Add a styled heading"""
    h = doc.add_heading(text, level=level)
    return h


def add_body_text(doc, text):
    """Add body paragraph"""
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(10)
    return p


def add_bullet(doc, text):
    """Add bullet point"""
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(10)
    return p


def add_figure_box(doc, title, content_lines, border_color='1a365d', bg_color='F0F4F8'):
    """Add a figure-like box"""
    outer_table = doc.add_table(rows=1, cols=1)
    outer_table.style = 'Table Grid'
    outer_cell = outer_table.rows[0].cells[0]
    set_cell_shading(outer_cell, bg_color)

    p = outer_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(
        int(border_color[0:2], 16),
        int(border_color[2:4], 16),
        int(border_color[4:6], 16)
    )

    for line in content_lines:
        p = outer_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(9)

    return outer_table


def create_v300_report():
    """Create v3.0.0 technical report"""

    # Try to load existing document as base
    src_path = os.path.join(os.path.dirname(__file__),
                            'PQC_FHE_Technical_Report_v2.3.5_Enterprise.docx')

    if os.path.exists(src_path):
        doc = Document(src_path)
        print(f"Loaded existing report: {src_path}")
    else:
        doc = Document()
        print("Creating new document from scratch")

    # =====================================================================
    # Add v3.0.0 sections at the end (before References if possible)
    # =====================================================================

    doc.add_page_break()

    # =====================================================================
    # Section: v3.0.0 New Features Overview
    # =====================================================================
    add_heading(doc, 'v3.0.0 New Features', level=1)

    add_body_text(doc, (
        'Version 3.0.0 introduces three major new capabilities: '
        'Quantum Threat Simulation for assessing cryptographic vulnerability to quantum attacks, '
        'Security Scoring Framework for NIST IR 8547 compliance assessment, and '
        'MPC-HE 2-Party Private Inference for privacy-preserving multi-party computation. '
        'These additions transform the platform from a cryptographic toolkit into a comprehensive '
        'enterprise quantum-readiness assessment and migration platform.'
    ))

    add_styled_table(doc,
        ['Feature', 'Module', 'Key Capability'],
        [
            ['Quantum Threat Simulator', 'quantum_threat_simulator.py',
             'Shor/Grover resource estimation, threat timeline'],
            ['Security Scoring', 'security_scoring.py',
             'NIST IR 8547 compliance, 0-100 scoring'],
            ['MPC-HE Inference', 'mpc_he_inference.py',
             '2-party encrypted computation protocol'],
            ['Extended Benchmarks', 'benchmarks/__init__.py',
             'GPU vs CPU, quantum threat benchmarks'],
            ['New API Endpoints', 'api/server.py',
             '15 new endpoints (50 total)'],
            ['Web UI Tabs', 'web_ui/index.html',
             '3 new tabs (6 total)'],
        ],
        '2D8B4E'
    )
    doc.add_paragraph()

    # =====================================================================
    # Section: Quantum Threat Simulator
    # =====================================================================
    doc.add_page_break()
    add_heading(doc, 'Quantum Threat Simulator', level=1)

    add_body_text(doc, (
        "The Quantum Threat Simulator provides scientific resource estimation for "
        "quantum attacks against classical and post-quantum cryptographic algorithms. "
        "It implements models based on Gidney & Ekera (2021) for Shor's algorithm and "
        "Grassl et al. (2016) for Grover's algorithm, enabling organizations to assess "
        "the timeline and feasibility of quantum threats to their cryptographic infrastructure."
    ))

    # Shor's Algorithm
    add_heading(doc, "Shor's Algorithm Resource Estimation", level=2)

    add_body_text(doc, (
        "Shor's algorithm enables polynomial-time factoring of integers and discrete logarithm "
        "computation, threatening RSA, ECC, and Diffie-Hellman key exchanges. The simulator uses "
        "the Gidney-Ekera (2021) optimized model requiring 2n+1 logical qubits for n-bit RSA keys, "
        "with an error correction overhead of 1000 physical qubits per logical qubit."
    ))

    add_styled_table(doc,
        ['Algorithm', 'Key Size', 'Logical Qubits', 'Physical Qubits', 'Threat Level', 'Est. Threat Year'],
        [
            ['RSA', '1024-bit', '2,049', '2,049,000', 'Moderate', '~2046'],
            ['RSA', '2048-bit', '4,097', '4,097,000', 'Low', '~2048'],
            ['RSA', '3072-bit', '6,145', '6,145,000', 'Low', '~2050'],
            ['RSA', '4096-bit', '8,193', '8,193,000', 'Low', '~2050'],
            ['ECC', 'P-256', '626,768', '626,768,000', 'Low', '~2048'],
            ['ECC', 'P-384', '940,544', '940,544,000', 'Low', '~2050'],
        ],
        'C53030'
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Table: Shor's Algorithm Quantum Resource Estimation")
    run.font.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Grover's Algorithm
    add_heading(doc, "Grover's Algorithm Resource Estimation", level=2)

    add_body_text(doc, (
        "Grover's algorithm provides a quadratic speedup for unstructured search, "
        "effectively halving the security of symmetric cryptographic algorithms. "
        "AES-128 provides only 64-bit post-quantum security, making AES-256 the "
        "minimum recommended key size for quantum resistance."
    ))

    add_styled_table(doc,
        ['Algorithm', 'Key Size', 'Logical Qubits', 'Post-Quantum Security', 'Threat Level'],
        [
            ['AES', '128-bit', '2,953', '64 bits', 'Moderate'],
            ['AES', '192-bit', '4,449', '96 bits', 'Low'],
            ['AES', '256-bit', '6,681', '128 bits', 'Low'],
            ['SHA', '256-bit', '~3,000', '128 bits', 'Low'],
            ['SHA', '384-bit', '~4,500', '192 bits', 'Low'],
            ['SHA', '512-bit', '~6,000', '256 bits', 'Low'],
        ],
        'B7791A'
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Table: Grover's Algorithm Quantum Resource Estimation")
    run.font.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Quantum Threat Timeline
    add_heading(doc, 'Quantum Threat Timeline', level=2)

    add_body_text(doc, (
        "The threat timeline projects when quantum computers may achieve sufficient "
        "qubit counts to execute attacks, based on three QPU growth models: "
        "Conservative (linear growth), Moderate (polynomial), and Aggressive (exponential). "
        "Under the moderate model, RSA-2048 faces a credible threat by approximately 2048."
    ))

    add_figure_box(doc,
        'Quantum Threat Timeline (Moderate Growth Model)',
        [
            '',
            'RSA-1024: Threat ~2046 | RSA-2048: Threat ~2048 | RSA-4096: Threat ~2050',
            'ECC P-256: Threat ~2048 | ECC P-384: Threat ~2050',
            '',
            'AES-128: Moderate risk (64-bit PQ) | AES-256: Low risk (128-bit PQ)',
            '',
            'ML-KEM / ML-DSA: RESISTANT (lattice-based, no known quantum speedup)',
            '',
        ],
        'C53030', 'FFF5F5'
    )

    # Classical vs PQC
    add_heading(doc, 'Classical vs PQC Comparison', level=2)

    add_styled_table(doc,
        ['Category', 'Algorithm', 'Quantum Vulnerable', 'Recommended PQC Alternative'],
        [
            ['Key Exchange', 'RSA-2048', 'Yes (Shor)', 'ML-KEM-768'],
            ['Key Exchange', 'ECDH P-256', 'Yes (Shor)', 'ML-KEM-768'],
            ['Key Exchange', 'X25519', 'Yes (Shor)', 'ML-KEM-768 Hybrid'],
            ['Signatures', 'RSA-2048', 'Yes (Shor)', 'ML-DSA-65'],
            ['Signatures', 'ECDSA P-256', 'Yes (Shor)', 'ML-DSA-65'],
            ['Symmetric', 'AES-128', 'Weakened (Grover)', 'AES-256'],
            ['Symmetric', 'AES-256', 'Resistant', 'AES-256 (sufficient)'],
            ['Hash', 'SHA-256', 'Resistant', 'SHA-256 (sufficient)'],
        ],
        '2B6CB0'
    )

    # =====================================================================
    # Section: Security Scoring Framework
    # =====================================================================
    doc.add_page_break()
    add_heading(doc, 'Security Scoring Framework', level=1)

    add_body_text(doc, (
        "The Security Scoring Framework provides quantitative assessment of an organization's "
        "post-quantum cryptography readiness, aligned with NIST IR 8547 transition guidelines. "
        "It evaluates cryptographic inventories across five weighted dimensions and generates "
        "actionable migration plans."
    ))

    # Scoring Methodology
    add_heading(doc, 'Scoring Methodology', level=2)

    add_body_text(doc, (
        "The overall security score (0-100) is computed as a weighted average of five sub-scores, "
        "each reflecting a critical aspect of quantum readiness:"
    ))

    add_styled_table(doc,
        ['Sub-Score', 'Weight', 'Description', 'Measurement'],
        [
            ['Algorithm Strength', '25%', 'Strength of deployed algorithms', 'Key sizes, quantum resistance'],
            ['PQC Readiness', '25%', 'Adoption of PQC algorithms', '% assets using ML-KEM/ML-DSA'],
            ['Compliance', '20%', 'Regulatory compliance', 'NIST/CNSA/FIPS check results'],
            ['Key Management', '15%', 'Key lifecycle practices', 'Rotation, storage, distribution'],
            ['Crypto Agility', '15%', 'Ability to transition', 'Abstraction, hybrid support'],
        ],
        '6B46C1'
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Table: Security Scoring Weight Distribution')
    run.font.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Risk Categories
    add_heading(doc, 'Risk Categories', level=2)

    add_styled_table(doc,
        ['Score Range', 'Category', 'Action Required'],
        [
            ['0-25', 'Critical', 'Immediate remediation required'],
            ['26-50', 'High', 'Urgent migration planning needed'],
            ['51-75', 'Moderate', 'Active migration recommended'],
            ['76-100', 'Low', 'Maintain and monitor'],
        ],
        'C53030'
    )

    # Sample Assessments
    add_heading(doc, 'Sample Assessment Results', level=2)

    add_body_text(doc, (
        "Three sample organizational profiles demonstrate the scoring framework in action:"
    ))

    add_styled_table(doc,
        ['Organization', 'Score', 'Category', 'Algo', 'PQC', 'Compliance', 'KeyMgmt', 'Agility'],
        [
            ['Enterprise (Tech)', '34.1', 'High', '65.4', '10.0', '38.9', '30.0', '20.0'],
            ['Financial (Bank)', '44.2', 'High', '72.0', '20.0', '45.0', '40.0', '30.0'],
            ['Government (Agency)', '51.0', 'Moderate', '75.0', '30.0', '55.0', '50.0', '40.0'],
        ],
        '2D8B4E'
    )

    # Compliance Standards
    add_heading(doc, 'Compliance Checks', level=2)

    add_styled_table(doc,
        ['Standard', 'Focus', 'Key Requirements'],
        [
            ['NIST IR 8547', 'PQC Migration', 'Crypto inventory, PQC deployment plan, hybrid mode'],
            ['NSA CNSA 2.0', 'National Security', 'ML-KEM-1024 by 2030, ML-DSA-87 required'],
            ['FIPS 140-3', 'Crypto Validation', 'CMVP validated modules, approved algorithms'],
            ['NIST SP 800-57', 'Key Management', 'Key lifecycle, rotation, secure storage'],
        ],
        '2B6CB0'
    )

    # Migration Plan
    add_heading(doc, 'Migration Plan Generation', level=2)

    add_body_text(doc, (
        "Based on the assessment score, the framework generates a 4-phase migration plan "
        "aligned with NIST IR 8547 recommendations:"
    ))

    add_styled_table(doc,
        ['Phase', 'Name', 'Timeline', 'Key Actions'],
        [
            ['1', 'Assessment', '2024-2025', 'Complete cryptographic inventory, identify vulnerable assets'],
            ['2', 'Hybrid Deployment', '2025-2027', 'Deploy hybrid classical+PQC for high-value systems'],
            ['3', 'PQC Primary', '2027-2030', 'Make PQC primary with classical fallback'],
            ['4', 'Full Migration', '2030-2035', 'Remove all classical algorithms, PQC only'],
        ],
        '2D8B4E'
    )

    # =====================================================================
    # Section: MPC-HE 2-Party Private Inference
    # =====================================================================
    doc.add_page_break()
    add_heading(doc, 'MPC-HE 2-Party Private Inference', level=1)

    add_body_text(doc, (
        "The Multi-Party Computation with Homomorphic Encryption (MPC-HE) module implements "
        "a 2-party private inference protocol using the DESILO FHE multiparty API. This enables "
        "ALICE (data owner) and BOB (model owner) to collaboratively compute on encrypted data "
        "without either party revealing their private inputs."
    ))

    # Protocol Overview
    add_heading(doc, 'Protocol Overview', level=2)

    add_styled_table(doc,
        ['Phase', 'Name', 'Operations', 'Participants'],
        [
            ['1', 'Key Setup', 'Generate common public/evaluation keys', 'ALICE + BOB'],
            ['2', 'Encryption', 'Encrypt input data with common public key', 'ALICE'],
            ['3', 'Computation', 'Apply model layers on encrypted data', 'BOB'],
            ['4', 'Decryption', 'Individual decrypt + multiparty combine', 'ALICE + BOB'],
        ],
        '805AD5'
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Table: MPC-HE 4-Phase Protocol')
    run.font.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # DESILO Multiparty API
    add_heading(doc, 'DESILO FHE Multiparty Integration', level=2)

    add_body_text(doc, (
        "The protocol leverages DESILO FHE's multiparty capabilities with the following key API calls:"
    ))

    add_styled_table(doc,
        ['Operation', 'DESILO API', 'Description'],
        [
            ['Engine Init', 'Engine(use_multiparty=True)', 'Enable multiparty mode'],
            ['Common Keys', 'create_multiparty_public_key()', 'Generate shared encryption key'],
            ['Encrypt', 'encrypt(data)', 'Encrypt with common public key'],
            ['Compute', 'add/multiply/rotate', 'Homomorphic operations on ciphertexts'],
            ['Individual Decrypt', 'individual_decrypt(ct)', 'Each party creates decryption share'],
            ['Final Decrypt', 'multiparty_decrypt(ct, shares)', 'Combine shares for plaintext'],
        ],
        '805AD5'
    )

    # Activation Functions
    add_heading(doc, 'Chebyshev Polynomial Activations', level=2)

    add_body_text(doc, (
        "For encrypted neural network inference, non-linear activation functions are approximated "
        "using Chebyshev polynomial expansions, maintaining FHE compatibility while preserving "
        "accuracy. The following activations are supported:"
    ))

    add_styled_table(doc,
        ['Activation', 'Polynomial Degree', 'Domain', 'Max Error'],
        [
            ['ReLU (smooth)', '3', '[-5, 5]', '< 0.05'],
            ['Sigmoid', '5', '[-5, 5]', '< 0.01'],
            ['GELU', '5', '[-5, 5]', '< 0.02'],
            ['Tanh', '5', '[-3, 3]', '< 0.01'],
            ['Swish', '5', '[-5, 5]', '< 0.02'],
            ['Exponential', '7', '[-5, 5]', '< 0.05'],
        ],
        'DD6B20'
    )

    # BFV Note
    add_heading(doc, 'BFV Scheme Support', level=2)

    add_body_text(doc, (
        "The BFV (Brakerski/Fan-Vercauteren) homomorphic encryption scheme is not natively "
        "supported by the DESILO FHE library, which focuses on the CKKS scheme for approximate "
        "arithmetic. For BFV-based integer arithmetic workloads, the platform provides guidance "
        "for integration with the HEonGPU C++ library (RhombusEnd2End_HEonGPU), which supports "
        "GPU-accelerated BFV/CKKS/TFHE operations. A NotImplementedError with integration "
        "instructions is raised when BFV mode is requested."
    ))

    # Demo Scenarios
    add_heading(doc, 'Demo Scenarios', level=2)

    add_styled_table(doc,
        ['Demo', 'Description', 'ALICE Input', 'BOB Model'],
        [
            ['Linear Regression', 'Encrypted linear prediction', 'Feature vector', 'Weight matrix + bias'],
            ['Classification', 'Encrypted binary classification', 'Feature vector', 'Logistic model'],
            ['Private Statistics', 'Encrypted statistical aggregation', 'Dataset A', 'Dataset B'],
        ],
        '2D8B4E'
    )

    # =====================================================================
    # Section: Extended Benchmarks
    # =====================================================================
    doc.add_page_break()
    add_heading(doc, 'Extended Benchmarks (v3.0.0)', level=1)

    add_body_text(doc, (
        "Version 3.0.0 extends the benchmark suite with quantum threat estimation, security "
        "scoring, MPC-HE protocol, and GPU vs CPU comparison benchmarks. "
        "Performance measurements target the NVIDIA RTX 6000 PRO Blackwell 96GB GPU with "
        "CUDA 13.0 for GPU-accelerated FHE operations."
    ))

    # Development Environment
    add_heading(doc, 'Development Environment', level=2)

    add_styled_table(doc,
        ['Component', 'Specification'],
        [
            ['OS', 'Alma Linux 9.7'],
            ['CPU', 'Intel Core i5-13600K'],
            ['RAM', '128GB DDR5 5200'],
            ['GPU', 'NVIDIA RTX 6000 PRO Blackwell 96GB'],
            ['Python', '3.12.11'],
            ['CUDA', '13.0'],
            ['liboqs', '0.15.0'],
            ['DESILO FHE', 'desilofhe-cu130 (GPU mode)'],
        ],
        '2B6CB0'
    )

    # Benchmark Categories
    add_heading(doc, 'Benchmark Categories', level=2)

    add_styled_table(doc,
        ['Category', 'Metrics', 'Iterations'],
        [
            ['Quantum Threat Estimation', 'Shor/Grover resource calculation time', '10'],
            ['Security Scoring', 'Assessment computation time per inventory type', '10'],
            ['MPC-HE Protocol', 'Key setup, encrypt, compute, decrypt phases', '5'],
            ['GPU vs CPU FHE', 'Encrypt/decrypt/add/multiply/bootstrap speedup', '100'],
            ['Security-Performance', 'Trade-off across 128/192/256-bit levels', '10'],
        ],
        'DD6B20'
    )

    # GPU Benchmark Results
    add_heading(doc, 'GPU vs CPU Performance (Target)', level=2)

    add_body_text(doc, (
        "Target performance benchmarks for RTX 6000 PRO Blackwell with desilofhe-cu130. "
        "Actual results depend on DESILO FHE GPU backend availability."
    ))

    add_styled_table(doc,
        ['Operation', 'CPU Time', 'GPU Time', 'Expected Speedup'],
        [
            ['Key Generation', '~2.5 s', '~0.5 s', '~5x'],
            ['Encrypt (8192 slots)', '~15 ms', '~2 ms', '~7x'],
            ['Decrypt', '~10 ms', '~1.5 ms', '~7x'],
            ['Add (CT+CT)', '~0.5 ms', '~0.05 ms', '~10x'],
            ['Multiply (CT*CT)', '~50 ms', '~5 ms', '~10x'],
            ['Bootstrap', '~15 s', '~1.5 s', '~10x'],
        ],
        '2D8B4E'
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Table: Expected GPU vs CPU Performance (RTX 6000 PRO Blackwell)')
    run.font.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # =====================================================================
    # Section: Updated API Reference
    # =====================================================================
    doc.add_page_break()
    add_heading(doc, 'New API Endpoints (v3.0.0)', level=1)

    add_body_text(doc, (
        "Version 3.0.0 adds 15 new REST API endpoints across four categories, "
        "bringing the total to 50 endpoints with automatic OpenAPI/Swagger documentation."
    ))

    # Quantum Threat API
    add_heading(doc, 'Quantum Threat Endpoints', level=2)

    add_styled_table(doc,
        ['Endpoint', 'Method', 'Description'],
        [
            ['/quantum/threat-assessment', 'POST', 'Full quantum threat assessment for specified algorithms'],
            ['/quantum/threat-timeline', 'GET', 'Quantum threat timeline with growth models'],
            ['/quantum/pqc-comparison', 'GET', 'Classical vs PQC vulnerability comparison'],
            ['/quantum/shor-simulation/{key_size}', 'GET', 'Shor algorithm resource estimation'],
            ['/quantum/grover-simulation/{key_size}', 'GET', 'Grover algorithm resource estimation'],
        ],
        'C53030'
    )

    # Security Scoring API
    add_heading(doc, 'Security Scoring Endpoints', level=2)

    add_styled_table(doc,
        ['Endpoint', 'Method', 'Description'],
        [
            ['/security/assess', 'POST', 'Security assessment with 0-100 scoring'],
            ['/security/compliance/{standard}', 'GET', 'Compliance check (nist_ir_8547/cnsa_2_0/fips_140_3)'],
            ['/security/migration-plan', 'GET', 'PQC migration plan generation'],
            ['/security/inventory-templates', 'GET', 'Sample cryptographic inventories'],
        ],
        '6B46C1'
    )

    # MPC-HE API
    add_heading(doc, 'MPC-HE Endpoints', level=2)

    add_styled_table(doc,
        ['Endpoint', 'Method', 'Description'],
        [
            ['/mpc-he/protocol-info', 'GET', 'MPC-HE protocol information and phases'],
            ['/mpc-he/demo/{demo_type}', 'POST', 'Run MPC-HE demo (linear_regression/classification/statistics)'],
        ],
        '805AD5'
    )

    # Extended Benchmarks API
    add_heading(doc, 'Extended Benchmark Endpoints', level=2)

    add_styled_table(doc,
        ['Endpoint', 'Method', 'Description'],
        [
            ['/benchmarks/quantum-threat', 'POST', 'Quantum threat estimation benchmark'],
            ['/benchmarks/security-scoring', 'POST', 'Security scoring benchmark'],
            ['/benchmarks/extended', 'POST', 'Run all extended benchmarks'],
        ],
        'DD6B20'
    )

    # =====================================================================
    # Section: Updated Roadmap
    # =====================================================================
    doc.add_page_break()
    add_heading(doc, 'Updated Development Roadmap', level=1)

    add_styled_table(doc,
        ['Version', 'Timeline', 'Major Features', 'Status'],
        [
            ['v2.3.5', '2025-12-30', 'Hybrid key exchange, K8s, monitoring, logging', 'Released'],
            ['v3.0.0', '2026-03-18', 'Quantum threat, security scoring, MPC-HE, GPU benchmarks', 'Released'],
            ['v3.1.0', 'Q2 2026', 'Hardware security module (HSM) integration', 'Planned'],
            ['v3.2.0', 'Q3 2026', 'Zero-knowledge proof support', 'Planned'],
            ['v3.3.0', 'Q4 2026', 'FIPS 140-3 CMVP validation', 'Planned'],
            ['v4.0.0', 'Q1 2027', 'Full TLS 1.3 PQC integration', 'Planned'],
        ],
        '2D8B4E'
    )

    # =====================================================================
    # Section: Updated References
    # =====================================================================
    add_heading(doc, 'Additional References (v3.0.0)', level=1)

    refs = [
        '[13] Gidney C, Ekera M. How to factor 2048 bit RSA integers in 8 hours using 20 million noisy qubits. Quantum 5, 433 (2021). DOI: 10.22331/q-2021-04-15-433',
        '[14] Grassl M, et al. Applying Grover\'s Algorithm to AES: Quantum Resource Estimates. PQCrypto 2016. DOI: 10.1007/978-3-319-29360-8_3',
        '[15] NIST IR 8547: Transition to Post-Quantum Cryptography Standards. https://csrc.nist.gov/pubs/ir/8547/final',
        '[16] NSA CNSA 2.0: Commercial National Security Algorithm Suite 2.0. https://media.defense.gov/2022/Sep/07/2003071836/-1/-1/0/CSI_CNSA_2.0_FAQ_.PDF',
        '[17] DESILO FHE Multiparty API. https://fhe.desilo.dev/latest/multiparty/',
        '[18] HEonGPU: GPU-accelerated BFV/CKKS/TFHE. https://github.com/RhombusEnd2End_HEonGPU',
    ]

    for ref in refs:
        p = doc.add_paragraph(ref)
        for run in p.runs:
            run.font.size = Pt(9)

    # =====================================================================
    # Footer
    # =====================================================================
    doc.add_paragraph()
    doc.add_paragraph()

    # Separator
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('_' * 70)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Document info
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Document Updated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('PQC-FHE Integration Platform v3.0.0 Enterprise Edition')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\u00A9 2026 - MIT License')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Save
    output_path = os.path.join(os.path.dirname(__file__),
                               'PQC_FHE_Technical_Report_v3.0.0_Enterprise.docx')
    doc.save(output_path)
    print(f"\nSaved: {output_path}")
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print(f"Total tables: {len(doc.tables)}")
    return output_path


if __name__ == '__main__':
    create_v300_report()
