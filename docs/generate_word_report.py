#!/usr/bin/env python3
"""
Generate Word Document for PQC-FHE Technical Report v2.3.4
Uses python-docx library
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_report():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    # Title style
    title_style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = 'Arial'
    title_style.font.size = Pt(28)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)
    
    # Heading1 style
    h1_style = doc.styles['Heading 1']
    h1_style.font.name = 'Arial'
    h1_style.font.size = Pt(16)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0x2c, 0x52, 0x82)
    
    # Heading2 style
    h2_style = doc.styles['Heading 2']
    h2_style.font.name = 'Arial'
    h2_style.font.size = Pt(13)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(0x31, 0x82, 0xce)
    
    # =========================================================================
    # TITLE PAGE
    # =========================================================================
    doc.add_paragraph()
    doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PQC-FHE Integration Platform")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Technical Report v2.3.4")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x4a, 0x55, 0x68)
    
    subtitle2 = doc.add_paragraph()
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle2.add_run("Post-Quantum Cryptography + Fully Homomorphic Encryption")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x4a, 0x55, 0x68)
    
    doc.add_paragraph()
    
    # Info table
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = 'Table Grid'
    info_data = [
        ('Version', '2.3.4'),
        ('Release Date', '2025-12-30'),
        ('PQC Standards', 'FIPS 203 (ML-KEM), FIPS 204 (ML-DSA)'),
        ('FHE Scheme', 'CKKS (DESILO Implementation)'),
        ('License', 'MIT'),
    ]
    for i, (label, value) in enumerate(info_data):
        row = info_table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x2c, 0x52, 0x82)
        row.cells[1].text = value
    
    doc.add_page_break()
    
    # =========================================================================
    # 1. EXECUTIVE SUMMARY
    # =========================================================================
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        "The PQC-FHE Integration Platform provides a production-ready framework combining "
        "Post-Quantum Cryptography (PQC) with Fully Homomorphic Encryption (FHE) for "
        "enterprise security applications. This platform addresses the emerging threat of "
        "quantum computers to current cryptographic systems while enabling privacy-preserving "
        "computation on sensitive data."
    )
    
    doc.add_heading('Key Features', level=2)
    features = [
        "NIST-standardized PQC algorithms (FIPS 203, FIPS 204)",
        "CKKS-based homomorphic encryption via DESILO FHE",
        "Real-time data integration from verified public sources",
        "REST API with comprehensive Swagger documentation",
        "Interactive Web UI with live demonstrations",
        "GPU acceleration support (CUDA 12.x/13.x)",
    ]
    for f in features:
        doc.add_paragraph(f, style='List Bullet')
    
    doc.add_heading('v2.3.4 Updates (2025-12-30)', level=2)
    doc.add_paragraph(
        "Version 2.3.4 introduces enhanced live data fetching capabilities with robust "
        "fallback mechanisms, fixes for numpy array handling in FHE demo endpoints, "
        "and improved multi-platform installation documentation."
    )
    
    doc.add_page_break()
    
    # =========================================================================
    # 2. MARKET CONTEXT
    # =========================================================================
    doc.add_heading('2. Market Context and Business Value', level=1)
    
    doc.add_heading('The Quantum Threat', level=2)
    doc.add_paragraph(
        "Quantum computers pose an existential threat to current public-key cryptography. "
        "Shor's algorithm can break RSA-2048 in polynomial time, and Grover's algorithm "
        "reduces symmetric key security by half. NIST estimates cryptographically-relevant "
        "quantum computers may emerge within 10-15 years."
    )
    
    doc.add_heading('Market Opportunity', level=2)
    market_table = doc.add_table(rows=4, cols=4)
    market_table.style = 'Table Grid'
    
    # Header row
    header_cells = market_table.rows[0].cells
    headers = ['Segment', '2024', '2034', 'CAGR']
    for i, h in enumerate(headers):
        header_cells[i].text = h
        header_cells[i].paragraphs[0].runs[0].font.bold = True
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(header_cells[i], '2c5282')
    
    market_data = [
        ('PQC Solutions', '$302M', '$30B', '58%'),
        ('FHE Market', '$200M', '$2.5B', '28%'),
        ('Quantum Security', '$1.2B', '$15B', '29%'),
    ]
    for i, row_data in enumerate(market_data):
        row = market_table.rows[i + 1]
        for j, val in enumerate(row_data):
            row.cells[j].text = val
    
    p = doc.add_paragraph()
    run = p.add_run("Source: Markets and Markets, Gartner (2024)")
    run.font.italic = True
    run.font.size = Pt(9)
    
    doc.add_page_break()
    
    # =========================================================================
    # 3. SYSTEM ARCHITECTURE
    # =========================================================================
    doc.add_heading('3. System Architecture', level=1)
    
    doc.add_heading('Component Overview', level=2)
    arch_table = doc.add_table(rows=7, cols=3)
    arch_table.style = 'Table Grid'
    
    # Header
    header_cells = arch_table.rows[0].cells
    for i, h in enumerate(['Layer', 'Component', 'Technology']):
        header_cells[i].text = h
        header_cells[i].paragraphs[0].runs[0].font.bold = True
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(header_cells[i], '3182ce')
    
    arch_data = [
        ('Presentation', 'Web UI', 'React + Tailwind CSS'),
        ('API', 'REST Server', 'FastAPI + Swagger UI'),
        ('Cryptography', 'PQC Manager', 'liboqs-python (ML-KEM, ML-DSA)'),
        ('Cryptography', 'FHE Engine', 'DESILO FHE (CKKS)'),
        ('Data', 'Live Fetcher', 'VitalDB, yfinance, Ethereum RPC'),
        ('Infrastructure', 'GPU Accel.', 'CUDA 12.x/13.x + cuQuantum'),
    ]
    for i, row_data in enumerate(arch_data):
        row = arch_table.rows[i + 1]
        for j, val in enumerate(row_data):
            row.cells[j].text = val
    
    doc.add_heading('Data Flow', level=2)
    flow_steps = [
        "Client sends request to REST API (FastAPI server on port 8000)",
        "API validates input and routes to appropriate handler",
        "For PQC operations: liboqs-python performs key generation/signing",
        "For FHE operations: DESILO engine encrypts/computes/decrypts",
        "Live data fetcher retrieves external data with automatic fallback",
        "Response returned to client with full audit trail",
    ]
    for i, step in enumerate(flow_steps, 1):
        doc.add_paragraph(f"{i}. {step}")
    
    doc.add_page_break()
    
    # =========================================================================
    # 4. PQC IMPLEMENTATION
    # =========================================================================
    doc.add_heading('4. Post-Quantum Cryptography Implementation', level=1)
    
    doc.add_heading('Key Encapsulation Mechanisms (FIPS 203)', level=2)
    kem_table = doc.add_table(rows=4, cols=5)
    kem_table.style = 'Table Grid'
    
    # Header
    header_cells = kem_table.rows[0].cells
    for i, h in enumerate(['Algorithm', 'Level', 'Public Key', 'Ciphertext', 'Use Case']):
        header_cells[i].text = h
        header_cells[i].paragraphs[0].runs[0].font.bold = True
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(header_cells[i], '48bb78')
    
    kem_data = [
        ('ML-KEM-512', '1', '800 B', '768 B', 'IoT/Embedded'),
        ('ML-KEM-768', '3', '1,184 B', '1,088 B', 'General Purpose'),
        ('ML-KEM-1024', '5', '1,568 B', '1,568 B', 'High Security'),
    ]
    for i, row_data in enumerate(kem_data):
        row = kem_table.rows[i + 1]
        for j, val in enumerate(row_data):
            row.cells[j].text = val
    
    doc.add_paragraph()
    
    doc.add_heading('Digital Signatures (FIPS 204)', level=2)
    sig_table = doc.add_table(rows=4, cols=5)
    sig_table.style = 'Table Grid'
    
    # Header
    header_cells = sig_table.rows[0].cells
    for i, h in enumerate(['Algorithm', 'Level', 'Public Key', 'Signature', 'Speed']):
        header_cells[i].text = h
        header_cells[i].paragraphs[0].runs[0].font.bold = True
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(header_cells[i], 'ed8936')
    
    sig_data = [
        ('ML-DSA-44', '2', '1,312 B', '2,420 B', 'Fastest'),
        ('ML-DSA-65', '3', '1,952 B', '3,309 B', 'Balanced'),
        ('ML-DSA-87', '5', '2,592 B', '4,627 B', 'Maximum'),
    ]
    for i, row_data in enumerate(sig_data):
        row = sig_table.rows[i + 1]
        for j, val in enumerate(row_data):
            row.cells[j].text = val
    
    doc.add_page_break()
    
    # =========================================================================
    # 5. FHE IMPLEMENTATION
    # =========================================================================
    doc.add_heading('5. Fully Homomorphic Encryption Implementation', level=1)
    
    doc.add_heading('CKKS Scheme Overview', level=2)
    doc.add_paragraph(
        "The platform uses the CKKS (Cheon-Kim-Kim-Song) homomorphic encryption scheme, "
        "which supports approximate arithmetic on encrypted real numbers. CKKS is ideal "
        "for machine learning and statistical analysis on encrypted data."
    )
    
    doc.add_heading('DESILO FHE Configuration', level=2)
    config_table = doc.add_table(rows=6, cols=3)
    config_table.style = 'Table Grid'
    
    # Header
    header_cells = config_table.rows[0].cells
    for i, h in enumerate(['Parameter', 'Value', 'Description']):
        header_cells[i].text = h
        header_cells[i].paragraphs[0].runs[0].font.bold = True
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(header_cells[i], '805ad5')
    
    config_data = [
        ('poly_degree', '16,384', 'Polynomial ring dimension'),
        ('coeff_mod_bit_sizes', '[60,40,40,40,60]', 'Coefficient modulus chain'),
        ('scale', '2^40', 'Encoding scale for precision'),
        ('max_mult_depth', '4', 'Maximum multiplicative depth'),
        ('slot_count', '8,192', 'Number of plaintext slots'),
    ]
    for i, row_data in enumerate(config_data):
        row = config_table.rows[i + 1]
        for j, val in enumerate(row_data):
            row.cells[j].text = val
    
    doc.add_page_break()
    
    # =========================================================================
    # 6. LIVE DATA SOURCES
    # =========================================================================
    doc.add_heading('6. Live Data Sources', level=1)
    doc.add_paragraph(
        "Version 2.3.4 provides robust live data fetching with automatic fallback to "
        "embedded sample data."
    )
    
    doc.add_heading('Healthcare: VitalDB', level=2)
    doc.add_paragraph("Dataset: VitalDB Open Dataset (6,388 surgical cases)", style='List Bullet')
    doc.add_paragraph("Method: vitaldb Python library", style='List Bullet')
    doc.add_paragraph("DOI: 10.1038/s41597-022-01411-5", style='List Bullet')
    doc.add_paragraph("License: CC BY-NC-SA 4.0", style='List Bullet')
    
    doc.add_heading('Finance: Yahoo Finance', level=2)
    doc.add_paragraph("Method: yfinance Python library", style='List Bullet')
    doc.add_paragraph("Symbols: AAPL, MSFT, GOOGL, AMZN, NVDA", style='List Bullet')
    
    doc.add_heading('Blockchain: Ethereum RPC Endpoints', level=2)
    rpc_table = doc.add_table(rows=4, cols=3)
    rpc_table.style = 'Table Grid'
    
    # Header
    header_cells = rpc_table.rows[0].cells
    for i, h in enumerate(['Priority', 'Endpoint', 'Provider']):
        header_cells[i].text = h
        header_cells[i].paragraphs[0].runs[0].font.bold = True
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(header_cells[i], '4a5568')
    
    rpc_data = [
        ('1', 'rpc.ankr.com/eth', 'Ankr (Primary)'),
        ('2', 'ethereum-rpc.publicnode.com', 'PublicNode'),
        ('3', 'cloudflare-eth.com', 'Cloudflare'),
    ]
    for i, row_data in enumerate(rpc_data):
        row = rpc_table.rows[i + 1]
        for j, val in enumerate(row_data):
            row.cells[j].text = val
    
    doc.add_page_break()
    
    # =========================================================================
    # 9. INSTALLATION GUIDE
    # =========================================================================
    doc.add_heading('9. Installation Guide (Multi-Platform)', level=1)
    
    doc.add_heading('Build Dependencies by Platform', level=2)
    
    p = doc.add_paragraph()
    run = p.add_run("Debian/Ubuntu:")
    run.font.bold = True
    p = doc.add_paragraph()
    run = p.add_run("sudo apt install -y cmake gcc g++ libssl-dev python3-dev git")
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    
    p = doc.add_paragraph()
    run = p.add_run("Fedora/RHEL/CentOS:")
    run.font.bold = True
    p = doc.add_paragraph()
    run = p.add_run("sudo dnf install -y cmake gcc gcc-c++ openssl-devel python3-devel git")
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    
    p = doc.add_paragraph()
    run = p.add_run("Arch Linux:")
    run.font.bold = True
    p = doc.add_paragraph()
    run = p.add_run("sudo pacman -S cmake gcc openssl python git")
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    
    p = doc.add_paragraph()
    run = p.add_run("macOS (Homebrew):")
    run.font.bold = True
    p = doc.add_paragraph()
    run = p.add_run("brew install cmake openssl@3 git")
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    
    doc.add_heading('Installing liboqs-python', level=2)
    doc.add_paragraph(
        "liboqs-python is NOT available via pip. It must be built from source. "
        "The recommended method auto-downloads and builds liboqs at runtime:"
    )
    
    p = doc.add_paragraph()
    run = p.add_run("Option A: Automatic Build (Recommended)")
    run.font.bold = True
    
    code_lines = [
        "git clone --depth=1 https://github.com/open-quantum-safe/liboqs-python",
        "cd liboqs-python && pip install . && cd .."
    ]
    for line in code_lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    
    p = doc.add_paragraph()
    run = p.add_run("Option B: Manual Build")
    run.font.bold = True
    
    manual_lines = [
        "# Build liboqs C library",
        "git clone --depth=1 https://github.com/open-quantum-safe/liboqs",
        "cmake -S liboqs -B liboqs/build -DBUILD_SHARED_LIBS=ON",
        "cmake --build liboqs/build --parallel 8",
        "sudo cmake --build liboqs/build --target install",
        "export LD_LIBRARY_PATH=$LD_LIBRARY_PATH:/usr/local/lib"
    ]
    for line in manual_lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    
    doc.add_page_break()
    
    # =========================================================================
    # REFERENCES
    # =========================================================================
    doc.add_heading('References', level=1)
    
    references = [
        "[1] NIST. FIPS 203: Module-Lattice-Based Key-Encapsulation Mechanism Standard. August 2024. https://csrc.nist.gov/pubs/fips/203/final",
        "[2] NIST. FIPS 204: Module-Lattice-Based Digital Signature Standard. August 2024. https://csrc.nist.gov/pubs/fips/204/final",
        "[3] NIST. FIPS 205: Stateless Hash-Based Digital Signature Standard. August 2024. https://csrc.nist.gov/pubs/fips/205/final",
        "[4] Cheon JH, et al. Homomorphic Encryption for Arithmetic of Approximate Numbers. ASIACRYPT 2017. DOI: 10.1007/978-3-319-70694-8_15",
        "[5] DESILO. DESILO FHE Library Documentation. https://fhe.desilo.dev/latest/",
        "[6] Lee HC, et al. VitalDB, a high-fidelity multi-parameter vital signs database. Scientific Data 9, 279 (2022). DOI: 10.1038/s41597-022-01411-5",
        "[7] Hebrail G, Berard A. Individual Household Electric Power Consumption Data Set. UCI ML Repository. DOI: 10.24432/C52G6F",
        "[8] Aroussi R. yfinance: Download market data from Yahoo! Finance API. https://github.com/ranaroussi/yfinance",
        "[9] Ethereum Foundation. Ethereum JSON-RPC API. https://ethereum.org/developers/docs/apis/json-rpc/",
        "[10] Open Quantum Safe Project. liboqs-python. https://github.com/open-quantum-safe/liboqs-python",
        "[11] FastAPI. Modern web framework for building APIs. https://fastapi.tiangolo.com/",
        "[12] American Heart Association. Understanding Blood Pressure Readings. https://www.heart.org/",
    ]
    
    for ref in references:
        p = doc.add_paragraph(ref)
        p.paragraph_format.space_after = Pt(6)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("---")
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Generated: 2025-12-30 | PQC-FHE Integration Platform v2.3.4")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    # Save document
    doc.save("PQC_FHE_Technical_Report_v2.3.4.docx")
    print("Word document generated: PQC_FHE_Technical_Report_v2.3.4.docx")

if __name__ == "__main__":
    create_report()
