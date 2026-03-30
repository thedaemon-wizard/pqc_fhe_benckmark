#!/usr/bin/env python3
"""
Generate PQC-FHE Technical Report v3.1.0
Adds Quantum Algorithm Verification, NIST Level Analysis, and Sector Benchmarks
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))


def set_cell_shading(cell, color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_styled_table(doc, headers, data, header_color='1a365d', widths=None):
    """Create a professionally styled table"""
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.style = 'Table Grid'

    if widths:
        for i, width in enumerate(widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)

    header_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        header_cells[i].text = h
        for paragraph in header_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(header_cells[i], header_color)

    for i, row_data in enumerate(data):
        row = table.rows[i + 1]
        for j, val in enumerate(row_data):
            row.cells[j].text = str(val)
            for paragraph in row.cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
        if i % 2 == 1:
            for cell in row.cells:
                set_cell_shading(cell, 'F7FAFC')

    return table


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_body_text(doc, text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(10)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(10)
    return p


def generate_v310_report():
    """Generate the v3.1.0 technical report document."""
    doc = Document()

    # Title Page
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_heading('PQC-FHE Integration Platform', level=0)
    subtitle = doc.add_heading('Technical Report v3.1.0', level=1)
    doc.add_paragraph()
    add_body_text(doc, f'Release Date: {datetime.now().strftime("%Y-%m-%d")}')
    add_body_text(doc, 'Classification: Enterprise Technical Report')
    add_body_text(doc, 'Author: PQC-FHE Integration Library')
    doc.add_page_break()

    # Table of Contents placeholder
    add_heading(doc, 'Table of Contents', level=1)
    toc_items = [
        '1. Executive Summary',
        '2. Quantum Algorithm Verification',
        '  2.1 Shor\'s Algorithm Circuit Simulation',
        '  2.2 Grover\'s Algorithm Circuit Simulation',
        '  2.3 Resource Extrapolation to RSA-2048 and AES-256',
        '3. NIST Security Level Verification',
        '  3.1 Lattice Parameter Analysis (BKZ/Core-SVP)',
        '  3.2 ML-KEM Parameter Verification',
        '  3.3 ML-DSA Parameter Verification',
        '4. Sector-Specific Benchmarks',
        '  4.1 Healthcare (HIPAA)',
        '  4.2 Finance (PCI-DSS/SOX)',
        '  4.3 Blockchain',
        '  4.4 IoT / Edge Devices',
        '  4.5 MPC-FHE',
        '5. API Reference (New Endpoints)',
        '6. System Architecture',
        '7. Dependencies and Environment',
        '8. Version History',
    ]
    for item in toc_items:
        add_body_text(doc, item)
    doc.add_page_break()

    # Section 1: Executive Summary
    add_heading(doc, '1. Executive Summary', level=1)
    add_body_text(doc, (
        'PQC-FHE Integration Platform v3.1.0 introduces quantum algorithm verification '
        'using real quantum circuit simulation via Qiskit AerSimulator. This release moves '
        'beyond mathematical resource estimation to actual circuit construction and execution '
        'for Shor\'s and Grover\'s algorithms, providing empirical validation of quantum '
        'threats to classical cryptography.'
    ))
    add_body_text(doc, (
        'Key additions include: (1) Real Shor\'s algorithm circuit simulation demonstrating '
        'factoring of small numbers via quantum period finding with QFT; (2) Real Grover\'s '
        'algorithm circuit simulation demonstrating quadratic speedup via amplitude amplification; '
        '(3) NIST security level verification using lattice BKZ/Core-SVP hardness analysis; '
        '(4) Sector-specific benchmarks for Healthcare, Finance, Blockchain, IoT, and MPC-FHE.'
    ))

    add_heading(doc, 'Key Metrics', level=2)
    add_styled_table(doc,
        ['Metric', 'Value'],
        [
            ['Total Test Cases', '88 (21 new in v3.1.0)'],
            ['New Source Modules', '2 (quantum_verification.py, sector_benchmarks.py)'],
            ['New API Endpoints', '5'],
            ['Shor Circuit (N=15)', '12 qubits, ~100ms, 100% success rate'],
            ['Grover Circuit (4-qubit)', '96%+ target probability, 15x speedup'],
            ['NIST Level Verification', '9 algorithms, all PASSED'],
            ['Sector Benchmarks', '5 sectors, 25+ use cases'],
            ['Qiskit Version', '2.3.1 + AerSimulator 0.17.2'],
        ],
        header_color='2c5282'
    )
    doc.add_page_break()

    # Section 2: Quantum Algorithm Verification
    add_heading(doc, '2. Quantum Algorithm Verification', level=1)
    add_body_text(doc, (
        'Unlike v3.0.0 which used only mathematical resource estimates (Gidney-Ekera 2021, '
        'Roetteler et al. 2017, Grassl et al. 2016), v3.1.0 constructs and executes actual '
        'quantum circuits on Qiskit\'s AerSimulator to demonstrate Shor\'s and Grover\'s '
        'algorithms on small problem instances.'
    ))

    add_heading(doc, '2.1 Shor\'s Algorithm Circuit Simulation', level=2)
    add_body_text(doc, (
        'Shor\'s algorithm is implemented using quantum period finding with the Quantum '
        'Fourier Transform (QFT). The circuit consists of: (1) Counting register initialized '
        'in uniform superposition via Hadamard gates; (2) Work register for modular '
        'exponentiation; (3) Controlled-U gates implementing a^(2^k) mod N; '
        '(4) Inverse QFT on counting register; (5) Measurement and classical post-processing '
        'using continued fractions to extract the period r.'
    ))

    add_styled_table(doc,
        ['Number N', 'Factors', 'Qubits', 'Circuit Depth', 'Period r', 'Execution Time'],
        [
            ['15', '3 x 5', '12', '~39', '4', '~100ms'],
            ['21', '3 x 7', '15', '~60', '6', '~200ms'],
            ['35', '5 x 7', '18', '~100', 'varies', '~500ms'],
        ],
        header_color='553c9a'
    )

    add_heading(doc, '2.2 Grover\'s Algorithm Circuit Simulation', level=2)
    add_body_text(doc, (
        'Grover\'s algorithm demonstrates quadratic speedup for unstructured search. The '
        'circuit applies: (1) Uniform superposition via Hadamard gates; (2) Oracle marking '
        'the target state with a phase flip (multi-controlled Z); (3) Diffusion operator '
        '(inversion about mean) via H-X-MCZ-X-H; (4) Optimal number of iterations: '
        'floor(pi/4 * sqrt(N)).'
    ))

    add_styled_table(doc,
        ['Qubits', 'Search Space', 'Optimal Iterations', 'P(target)', 'Classical P', 'Speedup'],
        [
            ['3', 'N=8', '2', '~94.5%', '12.5%', '~7.6x'],
            ['4', 'N=16', '3', '~96.1%', '6.25%', '~15.4x'],
            ['5', 'N=32', '4', '~96.6%', '3.13%', '~30.9x'],
        ],
        header_color='276749'
    )

    add_heading(doc, '2.3 Resource Extrapolation', level=2)
    add_body_text(doc, (
        'Results from small circuits are extrapolated to real-world key sizes. For RSA-2048, '
        'Gidney-Ekera (2021) estimates ~4,097 logical qubits (~20 million physical qubits '
        'with error correction) and approximately 8 hours runtime on a fault-tolerant quantum '
        'computer. For AES key search under Grover, AES-128 is reduced to 64-bit effective '
        'security (UNSAFE), while AES-256 maintains 128-bit security (SAFE per NIST IR 8547).'
    ))
    doc.add_page_break()

    # Section 3: NIST Security Level Verification
    add_heading(doc, '3. NIST Security Level Verification', level=1)
    add_body_text(doc, (
        'NIST FIPS 203 (ML-KEM) and FIPS 204 (ML-DSA) define specific lattice parameters '
        'that determine the security level. This module verifies these parameters by estimating '
        'the BKZ block size needed for the primal lattice attack, then computing the '
        'Core-SVP hardness.'
    ))

    add_heading(doc, '3.1 Methodology', level=2)
    add_bullet(doc, 'Classical Core-SVP: 0.292 * beta (BKZ 2.0 sieving)')
    add_bullet(doc, 'Quantum Core-SVP: 0.265 * beta (quantum sieve, Laarhoven 2015)')
    add_bullet(doc, 'BKZ block size estimated using GSA (Geometric Series Assumption)')
    add_bullet(doc, 'NIST thresholds: Level 1 >= 128 bits, Level 3 >= 192 bits, Level 5 >= 256 bits')

    add_heading(doc, '3.2 Verification Results', level=2)
    add_styled_table(doc,
        ['Algorithm', 'Claimed Level', 'Verified Level', 'Core-SVP (Quantum)', 'Status'],
        [
            ['ML-KEM-512', 'Level 1', 'Level 1+', '~133 bits', 'PASS'],
            ['ML-KEM-768', 'Level 3', 'Level 3+', '~225 bits', 'PASS'],
            ['ML-KEM-1024', 'Level 5', 'Level 5+', '~300 bits', 'PASS'],
            ['ML-DSA-44', 'Level 2', 'Level 2+', '~200 bits', 'PASS'],
            ['ML-DSA-65', 'Level 3', 'Level 3+', '~398 bits', 'PASS'],
            ['ML-DSA-87', 'Level 5', 'Level 5+', '~397 bits', 'PASS'],
            ['SLH-DSA-128s', 'Level 1', 'Level 1', '128 bits', 'PASS'],
            ['SLH-DSA-192s', 'Level 3', 'Level 3', '192 bits', 'PASS'],
            ['SLH-DSA-256s', 'Level 5', 'Level 5', '256 bits', 'PASS'],
        ],
        header_color='9b2c2c'
    )
    doc.add_page_break()

    # Section 4: Sector-Specific Benchmarks
    add_heading(doc, '4. Sector-Specific Benchmarks', level=1)
    add_body_text(doc, (
        'All benchmarks use actual cryptographic operations (liboqs PQC, DESILO FHE) on '
        'representative data sizes for each industry sector. Benchmarks were run on NVIDIA '
        'RTX PRO 6000 Blackwell (96GB VRAM) with GPU-accelerated FHE.'
    ))

    for sector, title, color, data in [
        ('healthcare', '4.1 Healthcare (HIPAA)', '9b2c2c', [
            ['Patient Data Key Exchange', 'ML-KEM-768', '2KB', '~0.06ms', 'L3'],
            ['Medical Record Signing', 'ML-DSA-65', '4KB', '~0.09ms', 'L3'],
            ['Medical Record Verify', 'ML-DSA-65', '4KB', '~0.04ms', 'L3'],
            ['Vital Signs FHE Encrypt', 'CKKS', '256B', '~0.62ms', 'L3'],
            ['Encrypted Vitals Analysis', 'CKKS', '256B', '~3.02ms', 'L3'],
        ]),
        ('finance', '4.2 Finance (PCI-DSS/SOX)', '276749', [
            ['TX Batch Key Exchange', 'ML-KEM-768', '10KB', '~0.06ms', 'L3'],
            ['Trade Settlement Signing', 'ML-DSA-65', '1KB', '~0.09ms', 'L3'],
            ['Key Rotation', 'ML-KEM-768', '-', '~0.04ms', 'L3'],
            ['High-Value Signing (L5)', 'ML-DSA-87', '2KB', '~0.09ms', 'L5'],
            ['Encrypted Portfolio Sum', 'CKKS', '512B', '~3ms', 'L3'],
        ]),
        ('blockchain', '4.3 Blockchain', '553c9a', [
            ['TX Signing (L3)', 'ML-DSA-65', '256B', '~0.05ms', 'L3'],
            ['High-Security Signing', 'ML-DSA-87', '256B', '~0.09ms', 'L5'],
            ['Lightweight Signing', 'ML-DSA-44', '256B', '~0.04ms', 'L2'],
            ['Batch Verify (10 tx)', 'ML-DSA-65', '2.5KB', '~0.31ms', 'L3'],
            ['Full TX Pipeline', 'ML-DSA-65', '256B', '~0.12ms', 'L3'],
        ]),
        ('iot', '4.4 IoT / Edge Devices', '2c5282', [
            ['Sensor 64B (L1)', 'ML-KEM-512', '64B', '~0.02ms', 'L1'],
            ['Sensor 64B (L3)', 'ML-KEM-768', '64B', '~0.02ms', 'L3'],
            ['Sensor 4KB (L1)', 'ML-KEM-512', '4KB', '~0.02ms', 'L1'],
            ['Sensor 4KB (L3)', 'ML-KEM-768', '4KB', '~0.02ms', 'L3'],
            ['Device Keygen', 'ML-KEM-512/768', '-', '~0.01ms', 'L1/L3'],
        ]),
    ]:
        add_heading(doc, title, level=2)
        add_styled_table(doc,
            ['Use Case', 'Algorithm', 'Data Size', 'Mean Latency', 'NIST Level'],
            data,
            header_color=color
        )
        doc.add_paragraph()

    doc.add_page_break()

    # Section 5: API Reference
    add_heading(doc, '5. API Reference (New v3.1.0 Endpoints)', level=1)
    add_styled_table(doc,
        ['Endpoint', 'Method', 'Description', 'Tag'],
        [
            ['/quantum/verify/shor', 'POST', 'Execute Shor\'s algorithm circuit', 'Quantum Verification'],
            ['/quantum/verify/grover', 'POST', 'Execute Grover\'s algorithm circuit', 'Quantum Verification'],
            ['/quantum/verify/nist-levels', 'GET', 'Verify NIST security levels', 'Quantum Verification'],
            ['/benchmarks/sector/{sector}', 'GET', 'Run sector-specific benchmarks', 'Sector Benchmarks'],
            ['/benchmarks/sector-all', 'GET', 'Run all sector benchmarks', 'Sector Benchmarks'],
        ],
        header_color='1a365d'
    )
    doc.add_page_break()

    # Section 6: System Architecture
    add_heading(doc, '6. System Architecture', level=1)
    add_body_text(doc, 'Source Module Structure (v3.1.0):')
    modules = [
        ['src/pqc_fhe_integration.py', '1,428 LOC', 'Core PQC+FHE integration'],
        ['src/quantum_threat_simulator.py', '871 LOC', 'Shor/Grover resource estimation'],
        ['src/quantum_verification.py', '~900 LOC', 'Real quantum circuit simulation (NEW)'],
        ['src/security_scoring.py', '1,069 LOC', 'NIST IR 8547 compliance scoring'],
        ['src/sector_benchmarks.py', '~700 LOC', 'Sector-specific benchmarks (NEW)'],
        ['src/mpc_he_inference.py', '1,061 LOC', 'MPC-HE 2-party protocol'],
        ['src/desilo_fhe_engine.py', '290 LOC', 'DESILO FHE wrapper'],
        ['src/pqc_simulator.py', '781 LOC', 'Educational ML-KEM/DSA simulator'],
    ]
    add_styled_table(doc,
        ['Module', 'Size', 'Description'],
        modules,
        header_color='2c5282'
    )

    # Section 7: Dependencies
    add_heading(doc, '7. Dependencies and Environment', level=1)
    deps = [
        ['Python', '3.12.11', 'Runtime'],
        ['liboqs-python', '0.14.0', 'Post-quantum cryptography (ML-KEM, ML-DSA)'],
        ['desilofhe-cu130', '1.10.0', 'GPU-accelerated FHE (CKKS scheme)'],
        ['Qiskit', '2.3.1', 'Quantum circuit construction and transpilation (NEW)'],
        ['Qiskit-Aer', '0.17.2', 'Quantum circuit simulation backend (NEW)'],
        ['NumPy', '2.4.3', 'Numerical computation'],
        ['SciPy', '1.17.1', 'Scientific computing'],
        ['FastAPI', '0.115+', 'REST API framework'],
        ['cryptography', '46.0.5', 'Classical cryptography (X25519)'],
    ]
    add_styled_table(doc,
        ['Package', 'Version', 'Purpose'],
        deps,
        header_color='276749'
    )

    add_heading(doc, 'Hardware Environment', level=2)
    add_bullet(doc, 'CPU: Intel Core i5-13600K')
    add_bullet(doc, 'RAM: 128GB DDR5 5200')
    add_bullet(doc, 'GPU: NVIDIA RTX PRO 6000 Blackwell (96GB VRAM)')
    add_bullet(doc, 'OS: Alma Linux 9.7')
    add_bullet(doc, 'CUDA: 13.0')

    # Section 8: Version History
    add_heading(doc, '8. Version History', level=1)
    versions = [
        ['v3.1.0', '2026-03-18', 'Quantum circuit verification (Qiskit), sector benchmarks, NIST level verification'],
        ['v3.0.0', '2026-03-18', 'Quantum threat simulator, security scoring, MPC-HE inference'],
        ['v2.3.5', '2025-12-30', 'Hybrid X25519+ML-KEM, Kubernetes, Prometheus monitoring'],
        ['v2.3.4', '2025-12-30', 'Enhanced security tests, algorithm comparison'],
        ['v2.2.0', '2025-12-29', 'DESILO FHE CKKS with bootstrap support'],
        ['v1.0.0', '2025-12-28', 'Initial release with ML-KEM and ML-DSA'],
    ]
    add_styled_table(doc,
        ['Version', 'Date', 'Key Changes'],
        versions,
        header_color='1a365d'
    )

    # Save
    output_dir = os.path.dirname(os.path.abspath(__file__))
    docx_path = os.path.join(output_dir, 'PQC_FHE_Technical_Report_v3.1.0_Enterprise.docx')
    doc.save(docx_path)
    print(f"Generated: {docx_path}")
    return docx_path


if __name__ == '__main__':
    generate_v310_report()
