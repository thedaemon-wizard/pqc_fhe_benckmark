#!/usr/bin/env python3
"""
PQC-FHE Integration Platform - Comprehensive Word Document v2.3.5
Enterprise-Grade Documentation - Full Version Matching PDF
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml
from datetime import datetime


def set_cell_shading(cell, color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_styled_table(doc, headers, data, header_color, widths=None):
    """Create a professionally styled table"""
    table = doc.add_table(rows=len(data)+1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Set column widths if provided
    if widths:
        for i, width in enumerate(widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)
    
    # Header row
    header_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        header_cells[i].text = h
        for paragraph in header_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(header_cells[i], header_color)
    
    # Data rows
    for i, row_data in enumerate(data):
        row = table.rows[i + 1]
        for j, val in enumerate(row_data):
            row.cells[j].text = str(val)
            for paragraph in row.cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
        # Alternating row colors
        if i % 2 == 1:
            for cell in row.cells:
                set_cell_shading(cell, 'F7FAFC')
    
    return table


def add_figure_box(doc, title, content_lines, border_color, bg_color):
    """Add a figure-like box with border"""
    # Create outer table for border effect
    outer_table = doc.add_table(rows=1, cols=1)
    outer_table.style = 'Table Grid'
    outer_cell = outer_table.rows[0].cells[0]
    set_cell_shading(outer_cell, bg_color)
    
    # Title
    p = outer_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(int(border_color[0:2], 16), int(border_color[2:4], 16), int(border_color[4:6], 16))
    
    # Content lines
    for line in content_lines:
        p = outer_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(9)
    
    return outer_table


def add_diagram_table(doc, title, rows_data, colors):
    """Create a diagram-like visualization using tables"""
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)
    
    doc.add_paragraph()
    
    # Create visual representation
    table = doc.add_table(rows=len(rows_data), cols=1)
    table.style = 'Table Grid'
    
    for i, (text, color) in enumerate(zip(rows_data, colors)):
        cell = table.rows[i].cells[0]
        cell.text = text
        set_cell_shading(cell, color)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                # White text for dark backgrounds
                if color in ['3182CE', '38A169', '805AD5', 'DD6B20', 'E53E3E', '667EEA']:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    
    return table


def add_caption(doc, text):
    """Add figure/table caption"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)


def create_report():
    doc = Document()
    
    # Configure styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    style.paragraph_format.space_after = Pt(6)
    
    h1_style = doc.styles['Heading 1']
    h1_style.font.name = 'Calibri'
    h1_style.font.size = Pt(18)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)
    h1_style.paragraph_format.space_before = Pt(24)
    h1_style.paragraph_format.space_after = Pt(12)
    
    h2_style = doc.styles['Heading 2']
    h2_style.font.name = 'Calibri'
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(0x2c, 0x52, 0x82)
    h2_style.paragraph_format.space_before = Pt(18)
    h2_style.paragraph_format.space_after = Pt(8)
    
    h3_style = doc.styles['Heading 3']
    h3_style.font.name = 'Calibri'
    h3_style.font.size = Pt(12)
    h3_style.font.bold = True
    h3_style.font.color.rgb = RGBColor(0x31, 0x82, 0xce)
    h3_style.paragraph_format.space_before = Pt(12)
    h3_style.paragraph_format.space_after = Pt(6)
    
    # =========================================================================
    # TITLE PAGE
    # =========================================================================
    for _ in range(4):
        doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PQC-FHE Integration Platform")
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)
    
    doc.add_paragraph()
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Technical Report v2.3.5 Enterprise Edition")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x4a, 0x55, 0x68)
    
    doc.add_paragraph()
    
    subtitle2 = doc.add_paragraph()
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle2.add_run("Post-Quantum Cryptography + Fully Homomorphic Encryption")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
    
    subtitle3 = doc.add_paragraph()
    subtitle3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle3.add_run("with Kubernetes Deployment and Production Monitoring")
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x7e, 0xea)
    
    for _ in range(3):
        doc.add_paragraph()
    
    # Info table
    info_table = doc.add_table(rows=9, cols=2)
    info_table.style = 'Table Grid'
    info_data = [
        ('Version', '2.3.5 Enterprise'),
        ('Release Date', '2025-12-30'),
        ('PQC Standards', 'FIPS 203 (ML-KEM), FIPS 204 (ML-DSA), FIPS 205 (SLH-DSA)'),
        ('Hybrid Mode', 'X25519 + ML-KEM-768 (IETF draft-ietf-tls-ecdhe-mlkem)'),
        ('FHE Scheme', 'CKKS (DESILO Implementation)'),
        ('Deployment', 'Kubernetes Helm Chart v1.0.0'),
        ('Monitoring', 'Prometheus + Grafana + AlertManager'),
        ('Logging', 'RotatingFileHandler (10MB × 5 backups)'),
        ('License', 'MIT License'),
    ]
    for i, (label, value) in enumerate(info_data):
        row = info_table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x2c, 0x52, 0x82)
        row.cells[1].text = value
        set_cell_shading(row.cells[0], 'EBF8FF')
    
    doc.add_page_break()
    
    # =========================================================================
    # TABLE OF CONTENTS
    # =========================================================================
    doc.add_heading('Table of Contents', level=1)
    
    toc_items = [
        ("1.", "Executive Summary", "3"),
        ("2.", "System Architecture", "4"),
        ("3.", "Post-Quantum Cryptography Implementation", "6"),
        ("4.", "Hybrid X25519 + ML-KEM Migration Strategy", "9"),
        ("5.", "Fully Homomorphic Encryption Implementation", "12"),
        ("6.", "Kubernetes Deployment", "14"),
        ("7.", "Monitoring and Observability", "17"),
        ("8.", "Logging System", "20"),
        ("9.", "API Reference", "22"),
        ("10.", "Enterprise Use Cases", "25"),
        ("11.", "Security Analysis", "27"),
        ("12.", "Performance Benchmarks", "29"),
        ("13.", "Future Roadmap", "31"),
        ("", "Appendix A: Helm Chart Configuration", "32"),
        ("", "Appendix B: Prometheus Alerting Rules", "33"),
        ("", "References", "34"),
    ]
    
    toc_table = doc.add_table(rows=len(toc_items), cols=3)
    for i, (num, title, page) in enumerate(toc_items):
        row = toc_table.rows[i]
        row.cells[0].text = num
        if num:
            row.cells[0].paragraphs[0].runs[0].font.bold = True
            row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x31, 0x82, 0xce)
        row.cells[1].text = title
        row.cells[2].text = page
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_page_break()
    
    # =========================================================================
    # 1. EXECUTIVE SUMMARY
    # =========================================================================
    doc.add_heading('1. Executive Summary', level=1)
    
    p = doc.add_paragraph()
    p.add_run(
        "The PQC-FHE Integration Platform v2.3.5 Enterprise Edition represents a comprehensive, "
        "production-ready framework that combines Post-Quantum Cryptography (PQC) with Fully "
        "Homomorphic Encryption (FHE) for enterprise security applications. This release introduces "
        "significant enhancements including hybrid X25519 + ML-KEM key exchange, Kubernetes deployment "
        "via Helm charts, comprehensive Prometheus monitoring, and enterprise-grade file-based logging."
    )
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(
        "The platform addresses the critical security challenges posed by the emergence of quantum "
        "computing technology. With NIST's finalization of post-quantum cryptography standards in "
        "August 2024, organizations now have clear guidance for migrating to quantum-resistant "
        "algorithms. This platform provides the tools and infrastructure necessary to begin that "
        "migration immediately."
    )
    
    doc.add_heading('1.1 Key Capabilities', level=2)
    
    capabilities = [
        ("Post-Quantum Cryptography", "Full implementation of NIST-standardized algorithms including "
         "ML-KEM (FIPS 203) for key encapsulation and ML-DSA (FIPS 204) for digital signatures, "
         "providing quantum-resistant security for sensitive communications."),
        ("Hybrid Key Exchange", "Defense-in-depth security combining classical X25519 with ML-KEM-768 "
         "following IETF draft-ietf-tls-ecdhe-mlkem specification, protecting against both current "
         "and future quantum threats."),
        ("Homomorphic Encryption", "CKKS scheme implementation via DESILO FHE library enabling "
         "computation on encrypted data without decryption, supporting privacy-preserving analytics "
         "across healthcare, finance, and IoT domains."),
        ("Enterprise Deployment", "Production-ready Kubernetes Helm chart with horizontal pod "
         "autoscaling (2-10 replicas), GPU worker support, Redis caching, and comprehensive monitoring."),
        ("Observability", "Integrated Prometheus metrics exposure, pre-configured alerting rules, "
         "and Grafana dashboard support for operational visibility."),
        ("Logging", "Rotating file-based logging with separate streams for server operations, "
         "errors, and HTTP access, supporting compliance and debugging requirements."),
    ]
    
    for title, desc in capabilities:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f"{title}: ")
        run.font.bold = True
        p.add_run(desc)
    
    doc.add_heading('1.2 Target Audience', level=2)
    
    p = doc.add_paragraph()
    p.add_run(
        "This platform is designed for enterprise security architects, DevOps engineers, and "
        "software developers who need to implement quantum-resistant cryptographic solutions "
        "while maintaining operational efficiency. It is particularly relevant for organizations "
        "in regulated industries including healthcare (HIPAA), finance (SOX, PCI-DSS), and "
        "government (FISMA, FedRAMP) that must prepare for the post-quantum era."
    )
    
    doc.add_heading('1.3 Document Organization', level=2)
    
    p = doc.add_paragraph()
    p.add_run(
        "This technical report is organized into 13 main sections covering architecture, "
        "cryptographic implementations, deployment, monitoring, and operational considerations. "
        "Appendices provide detailed configuration references and alerting rules. All sections "
        "include practical examples and performance benchmarks to guide implementation decisions."
    )
    
    doc.add_page_break()
    
    # =========================================================================
    # 2. SYSTEM ARCHITECTURE
    # =========================================================================
    doc.add_heading('2. System Architecture', level=1)
    
    p = doc.add_paragraph()
    p.add_run(
        "The PQC-FHE platform employs a layered architecture designed for scalability, security, "
        "and operational excellence. Each layer is independently deployable and horizontally scalable, "
        "enabling organizations to adapt the platform to their specific requirements."
    )
    
    doc.add_heading('2.1 Architecture Overview', level=2)
    
    p = doc.add_paragraph()
    p.add_run(
        "The system architecture consists of five distinct layers, each responsible for specific "
        "functionality. This separation of concerns enables independent scaling, easier maintenance, "
        "and improved security through isolation."
    )
    
    doc.add_paragraph()
    
    # Architecture diagram as table
    arch_layers = [
        ("PRESENTATION LAYER", "React Web UI + Tailwind CSS", "3182CE"),
        ("API LAYER", "FastAPI REST Server + Swagger UI + Prometheus Metrics", "38A169"),
        ("CRYPTOGRAPHY LAYER", "PQC Manager (ML-KEM/ML-DSA) | Hybrid Manager (X25519+ML-KEM) | FHE Engine (CKKS)", "805AD5"),
        ("DATA LAYER", "Live Data Sources (VitalDB, Yahoo Finance, Ethereum) | Logging System", "DD6B20"),
        ("INFRASTRUCTURE LAYER", "Kubernetes + Helm | Docker | GPU (CUDA) | Prometheus + Grafana", "E53E3E"),
    ]
    
    arch_table = doc.add_table(rows=len(arch_layers), cols=2)
    arch_table.style = 'Table Grid'
    for i, (layer, components, color) in enumerate(arch_layers):
        row = arch_table.rows[i]
        row.cells[0].text = layer
        row.cells[1].text = components
        set_cell_shading(row.cells[0], color)
        set_cell_shading(row.cells[1], color)
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.bold = True if cell == row.cells[0] else False
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    
    add_caption(doc, "Figure 2.1: PQC-FHE System Architecture Overview")
    
    doc.add_heading('2.2 Layer Descriptions', level=2)
    
    doc.add_heading('2.2.1 Presentation Layer', level=3)
    p = doc.add_paragraph()
    p.add_run(
        "The presentation layer provides a modern, responsive web interface built with React and "
        "styled using Tailwind CSS. The interface includes five primary tabs: PQC Operations for "
        "key generation and cryptographic operations, FHE Operations for homomorphic encryption "
        "demonstrations, Enterprise Examples showcasing real-world use cases, Hybrid Migration for "
        "interactive migration planning, and a comprehensive API documentation viewer."
    )
    
    doc.add_heading('2.2.2 API Layer', level=3)
    p = doc.add_paragraph()
    p.add_run(
        "The API layer implements a RESTful interface using FastAPI, providing automatic OpenAPI "
        "(Swagger) documentation, request validation via Pydantic models, and CORS support for "
        "cross-origin requests. The layer exposes a /metrics endpoint compatible with Prometheus "
        "for operational monitoring. All endpoints support JSON request/response formats with "
        "comprehensive error handling and rate limiting capabilities."
    )
    
    doc.add_heading('2.2.3 Cryptography Layer', level=3)
    p = doc.add_paragraph()
    p.add_run(
        "The cryptography layer consists of three specialized managers: the PQC Manager handles "
        "all post-quantum operations using liboqs-python, the Hybrid Manager coordinates combined "
        "X25519 + ML-KEM operations following IETF standards, and the FHE Engine manages homomorphic "
        "encryption operations using the DESILO library's CKKS scheme implementation. Each manager "
        "maintains its own key store and provides thread-safe operations."
    )
    
    doc.add_heading('2.2.4 Data Layer', level=3)
    p = doc.add_paragraph()
    p.add_run(
        "The data layer provides real-time data integration from verified public sources including "
        "VitalDB for healthcare vital signs, Yahoo Finance for market data, and Ethereum RPC for "
        "blockchain transactions. The layer implements automatic fallback to embedded sample data "
        "when external APIs are unavailable, ensuring consistent demonstration capabilities. The "
        "logging system provides comprehensive audit trails for all operations."
    )
    
    doc.add_heading('2.2.5 Infrastructure Layer', level=3)
    p = doc.add_paragraph()
    p.add_run(
        "The infrastructure layer supports multiple deployment models including Docker containers, "
        "Kubernetes orchestration via Helm charts, and optional GPU acceleration using CUDA 12.x/13.x. "
        "Redis provides distributed caching for session state and cryptographic key material, while "
        "Prometheus and Grafana deliver comprehensive monitoring and visualization capabilities."
    )
    
    doc.add_page_break()
    
    # Component table
    doc.add_heading('2.3 Component Summary', level=2)
    
    add_styled_table(doc,
        ['Component', 'Technology', 'Version', 'Purpose'],
        [
            ('Web UI', 'React + Tailwind CSS', '18.x / 3.x', 'User interface'),
            ('API Server', 'FastAPI + Uvicorn', '0.100+ / 0.25+', 'REST endpoints'),
            ('PQC Library', 'liboqs-python', '0.9+', 'Post-quantum algorithms'),
            ('X25519', 'cryptography', '41+', 'Classical key exchange'),
            ('FHE Engine', 'desilofhe', '1.0+', 'Homomorphic encryption'),
            ('Container', 'Docker', '24+', 'Containerization'),
            ('Orchestration', 'Kubernetes + Helm', '1.28+ / 3.13+', 'Deployment'),
            ('Monitoring', 'Prometheus + Grafana', '2.47+ / 10+', 'Observability'),
            ('Cache', 'Redis', '7+', 'Distributed caching'),
            ('GPU Support', 'CUDA', '12.x / 13.x', 'Acceleration (optional)'),
        ],
        '3182CE',
        [1.5, 2.0, 1.2, 2.0])
    
    doc.add_paragraph()
    add_caption(doc, "Table 2.1: Platform Component Summary")
    
    doc.add_page_break()
    
    # =========================================================================
    # 3. POST-QUANTUM CRYPTOGRAPHY IMPLEMENTATION
    # =========================================================================
    doc.add_heading('3. Post-Quantum Cryptography Implementation', level=1)
    
    p = doc.add_paragraph()
    p.add_run(
        "The platform implements NIST's finalized post-quantum cryptography standards, published on "
        "August 13, 2024. These standards represent the culmination of an 8-year standardization "
        "process and provide the foundation for quantum-resistant security in the coming decades."
    )
    
    doc.add_heading('3.1 The Quantum Threat', level=2)
    
    p = doc.add_paragraph()
    p.add_run(
        "Cryptographically-relevant quantum computers (CRQCs) pose an existential threat to current "
        "public-key cryptography. Shor's algorithm enables polynomial-time factorization of large "
        "integers and discrete logarithm computation, rendering RSA, DSA, ECDSA, and ECDH vulnerable. "
        "Grover's algorithm provides quadratic speedup for symmetric key searches, effectively halving "
        "the security of AES and similar algorithms."
    )
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(
        "The \"Harvest Now, Decrypt Later\" (HNDL) threat compounds this risk: adversaries can collect "
        "encrypted data today for decryption once quantum computers become available. This makes "
        "immediate migration critical for data requiring long-term confidentiality, including "
        "government secrets, medical records, financial data, and intellectual property."
    )
    
    doc.add_heading('3.2 NIST Post-Quantum Standards', level=2)
    
    p = doc.add_paragraph()
    p.add_run(
        "NIST finalized three post-quantum cryptography standards in August 2024:"
    )
    
    standards = [
        ("FIPS 203 (ML-KEM)", "Module-Lattice-Based Key Encapsulation Mechanism for secure key exchange"),
        ("FIPS 204 (ML-DSA)", "Module-Lattice-Based Digital Signature Algorithm for authentication"),
        ("FIPS 205 (SLH-DSA)", "Stateless Hash-Based Digital Signature Standard as backup option"),
    ]
    
    for std, desc in standards:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f"{std}: ")
        run.font.bold = True
        p.add_run(desc)
    
    doc.add_heading('3.3 Key Encapsulation Mechanisms (FIPS 203)', level=2)
    
    p = doc.add_paragraph()
    p.add_run(
        "ML-KEM (Module-Lattice-Based Key Encapsulation Mechanism) provides quantum-resistant key "
        "exchange based on the hardness of the Module Learning With Errors (MLWE) problem. The scheme "
        "offers three security levels with corresponding parameter sets, allowing organizations to "
        "choose the appropriate balance between security and performance."
    )
    
    doc.add_paragraph()
    
    add_styled_table(doc,
        ['Parameter', 'ML-KEM-512', 'ML-KEM-768', 'ML-KEM-1024'],
        [
            ('NIST Security Level', 'Level 1 (128-bit)', 'Level 3 (192-bit)', 'Level 5 (256-bit)'),
            ('Classical Equivalent', 'AES-128', 'AES-192', 'AES-256'),
            ('Public Key Size', '800 bytes', '1,184 bytes', '1,568 bytes'),
            ('Secret Key Size', '1,632 bytes', '2,400 bytes', '3,168 bytes'),
            ('Ciphertext Size', '768 bytes', '1,088 bytes', '1,568 bytes'),
            ('Shared Secret Size', '32 bytes', '32 bytes', '32 bytes'),
            ('Encapsulation Time', '~15 μs', '~20 μs', '~25 μs'),
            ('Decapsulation Time', '~15 μs', '~20 μs', '~30 μs'),
            ('Recommended Use', 'IoT, Embedded', 'General Purpose', 'High Security'),
        ],
        '48BB78',
        [1.6, 1.5, 1.5, 1.5])
    
    doc.add_paragraph()
    add_caption(doc, "Table 3.1: ML-KEM Parameter Comparison (FIPS 203)")
    
    doc.add_page_break()
    
    doc.add_heading('3.4 Digital Signature Algorithms (FIPS 204)', level=2)
    
    p = doc.add_paragraph()
    p.add_run(
        "ML-DSA (Module-Lattice-Based Digital Signature Algorithm) provides quantum-resistant digital "
        "signatures based on the Fiat-Shamir with Aborts paradigm over module lattices. The signature "
        "scheme offers deterministic signing with three security levels, ensuring consistent signature "
        "generation for the same message and key pair."
    )
    
    doc.add_paragraph()
    
    add_styled_table(doc,
        ['Parameter', 'ML-DSA-44', 'ML-DSA-65', 'ML-DSA-87'],
        [
            ('NIST Security Level', 'Level 2', 'Level 3', 'Level 5'),
            ('Public Key Size', '1,312 bytes', '1,952 bytes', '2,592 bytes'),
            ('Secret Key Size', '2,560 bytes', '4,032 bytes', '4,896 bytes'),
            ('Signature Size', '2,420 bytes', '3,309 bytes', '4,627 bytes'),
            ('Sign Time', '~100 μs', '~150 μs', '~200 μs'),
            ('Verify Time', '~50 μs', '~80 μs', '~100 μs'),
            ('Recommended Use', 'High Performance', 'Balanced', 'Maximum Security'),
        ],
        'ED8936',
        [1.6, 1.5, 1.5, 1.5])
    
    doc.add_paragraph()
    add_caption(doc, "Table 3.2: ML-DSA Parameter Comparison (FIPS 204)")
    
    doc.add_heading('3.5 Implementation Details', level=2)
    
    p = doc.add_paragraph()
    p.add_run(
        "The platform uses liboqs-python, the Python wrapper for the Open Quantum Safe liboqs library, "
        "which provides reference implementations of all NIST-standardized algorithms. Key features include:"
    )
    
    impl_features = [
        "Constant-time implementations to prevent timing side-channel attacks",
        "Thread-safe key generation and cryptographic operations",
        "Automatic algorithm negotiation based on peer capabilities",
        "Secure key storage with in-memory encryption",
        "Comprehensive error handling and validation",
    ]
    
    for feature in impl_features:
        doc.add_paragraph(f"• {feature}", style='List Bullet')
    
    doc.add_page_break()
    
    # =========================================================================
    # 4. HYBRID MIGRATION STRATEGY
    # =========================================================================
    doc.add_heading('4. Hybrid X25519 + ML-KEM Migration Strategy', level=1)
    
    p = doc.add_paragraph()
    p.add_run(
        "Hybrid cryptography combines classical and post-quantum algorithms to provide defense-in-depth "
        "security during the transition period. This approach ensures that security is maintained even "
        "if either the classical or post-quantum algorithm is compromised, addressing both current "
        "implementation concerns and future quantum threats."
    )
    
    doc.add_heading('4.1 Hybrid Key Exchange Protocol', level=2)
    
    p = doc.add_paragraph()
    p.add_run(
        "The hybrid key exchange protocol combines X25519 (classical elliptic curve Diffie-Hellman) "
        "with ML-KEM-768 (post-quantum lattice-based KEM) to derive a combined shared secret. The "
        "protocol follows the IETF draft-ietf-tls-ecdhe-mlkem specification for TLS 1.3 compatibility."
    )
    
    doc.add_paragraph()
    
    # Hybrid flow diagram as table
    flow_table = doc.add_table(rows=5, cols=3)
    flow_table.style = 'Table Grid'
    
    flow_data = [
        ("Alice (Sender)", "→", "Bob (Receiver)"),
        ("", "Bob's Public Keys (X25519 + ML-KEM)", "Has static keypair"),
        ("Generate ephemeral X25519", "←", ""),
        ("ML-KEM Encapsulate", "Ephemeral PK + Ciphertext →", ""),
        ("", "Both derive: SHA-256(X25519_SS || ML-KEM_SS)", "ML-KEM Decapsulate"),
    ]
    
    for i, (left, middle, right) in enumerate(flow_data):
        row = flow_table.rows[i]
        row.cells[0].text = left
        row.cells[1].text = middle
        row.cells[2].text = right
        if i == 0:
            for cell in row.cells:
                set_cell_shading(cell, '3182CE' if cell == row.cells[0] else ('38A169' if cell == row.cells[2] else 'E2E8F0'))
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) if cell != row.cells[1] else RGBColor(0x4a, 0x55, 0x68)
        elif i == 4:
            set_cell_shading(row.cells[1], '667EEA')
            for p in row.cells[1].paragraphs:
                for run in p.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    
    doc.add_paragraph()
    add_caption(doc, "Figure 4.1: Hybrid X25519 + ML-KEM Key Exchange Protocol")
    
    doc.add_heading('4.2 Why Hybrid Cryptography?', level=2)
    
    p = doc.add_paragraph()
    p.add_run(
        "The hybrid approach addresses several critical concerns in the post-quantum transition:"
    )
    
    hybrid_benefits = [
        ("Defense in Depth", "Security is maintained as long as at least one of the underlying "
         "algorithms remains secure. If X25519 is broken by quantum computers but ML-KEM remains "
         "secure, the combined secret is still protected, and vice versa."),
        ("HNDL Protection", "Data encrypted with hybrid key exchange is immediately protected "
         "against future quantum attacks, eliminating the \"harvest now, decrypt later\" vulnerability."),
        ("Implementation Redundancy", "Bugs or vulnerabilities discovered in one implementation "
         "do not immediately compromise security, providing time for patches while maintaining protection."),
        ("Regulatory Compliance", "Many standards bodies recommend or require hybrid approaches "
         "during the transition period, including guidance from NSA (CNSA 2.0) and BSI."),
        ("Smooth Migration Path", "Organizations can gradually transition from classical to "
         "post-quantum cryptography without breaking existing systems or requiring simultaneous upgrades."),
    ]
    
    for title, desc in hybrid_benefits:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f"{title}: ")
        run.font.bold = True
        p.add_run(desc)
    
    doc.add_page_break()
    
    doc.add_heading('4.3 IETF Compliance', level=2)
    
    p = doc.add_paragraph()
    p.add_run(
        "This implementation follows draft-ietf-tls-ecdhe-mlkem for TLS 1.3 hybrid key exchange. "
        "The combined shared secret is derived using concatenation followed by a key derivation function:"
    )
    
    doc.add_paragraph()
    code_p = doc.add_paragraph()
    code_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = code_p.add_run("Combined_SS = SHA-256(X25519_SharedSecret || ML-KEM_SharedSecret)")
    run.font.name = 'Courier New'
    run.font.size = Pt(11)
    run.font.bold = True
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(
        "This construction ensures that both algorithm contributions are incorporated into the final "
        "key material, and the output is a fixed 32-byte value suitable for symmetric key derivation."
    )
    
    doc.add_heading('4.4 Migration Timeline (NIST IR 8547)', level=2)
    
    p = doc.add_paragraph()
    p.add_run(
        "NIST IR 8547 provides guidance for transitioning to post-quantum cryptography. The recommended "
        "migration timeline spans approximately 10-15 years, with organizations encouraged to begin "
        "hybrid deployment immediately."
    )
    
    doc.add_paragraph()
    
    # Migration timeline as visual table
    timeline_table = doc.add_table(rows=2, cols=4)
    timeline_table.style = 'Table Grid'
    
    phases = [
        ("Phase 1\nAssessment\n2024-2025", "718096"),
        ("Phase 2\nHybrid ★\n2025-2027", "38A169"),
        ("Phase 3\nPQC Primary\n2027-2030", "3182CE"),
        ("Phase 4\nPQC Only\n2030-2035", "805AD5"),
    ]
    
    for i, (text, color) in enumerate(phases):
        cell = timeline_table.rows[0].cells[i]
        cell.text = text
        set_cell_shading(cell, color)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    
    actions = [
        "Inventory & planning",
        "Deploy hybrid mode",
        "PQC with fallback",
        "Full migration",
    ]
    
    for i, action in enumerate(actions):
        cell = timeline_table.rows[1].cells[i]
        cell.text = action
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(8)
    
    doc.add_paragraph()
    add_caption(doc, "Figure 4.2: NIST IR 8547 Migration Timeline (★ = Current Recommended Phase)")
    
    doc.add_paragraph()
    
    add_styled_table(doc,
        ['Phase', 'Timeline', 'Objective', 'Actions', 'Algorithms'],
        [
            ('1. Assessment', '2024-2025', 'Inventory', 'Identify all cryptographic assets', 'RSA, ECDSA, X25519'),
            ('2. Hybrid ★', '2025-2027', 'Deploy hybrid', 'Implement for high-value systems', 'X25519 + ML-KEM-768'),
            ('3. PQC Primary', '2027-2030', 'PQC first', 'Make PQC primary with fallback', 'ML-KEM-768, ML-DSA-65'),
            ('4. PQC Only', '2030-2035', 'Complete', 'Remove classical algorithms', 'ML-KEM-1024, ML-DSA-87'),
        ],
        '38A169',
        [1.0, 0.9, 0.9, 2.0, 1.5])
    
    doc.add_paragraph()
    add_caption(doc, "Table 4.1: PQC Migration Roadmap (NIST IR 8547)")
    
    doc.add_page_break()
    
    doc.add_heading('4.5 Algorithm Comparison', level=2)
    
    p = doc.add_paragraph()
    p.add_run(
        "The following table compares the three key exchange options: classical X25519, post-quantum "
        "ML-KEM-768, and the hybrid combination of both."
    )
    
    doc.add_paragraph()
    
    add_styled_table(doc,
        ['Property', 'X25519 (Classical)', 'ML-KEM-768 (PQC)', 'Hybrid (Combined)'],
        [
            ('Public Key Size', '32 bytes', '1,184 bytes', '1,216 bytes'),
            ('Ciphertext Size', '32 bytes', '1,088 bytes', '1,120 bytes'),
            ('Shared Secret', '32 bytes', '32 bytes', '32 bytes (SHA-256)'),
            ('Quantum Resistant', 'No', 'Yes', 'Yes'),
            ('Classical Secure', 'Yes', 'Assumed', 'Yes'),
            ('Key Generation', '~20 μs', '~25 μs', '~45 μs'),
            ('Encapsulation', '~20 μs', '~30 μs', '~50 μs'),
            ('Decapsulation', '~20 μs', '~25 μs', '~45 μs'),
            ('Total Latency', '~60 μs', '~80 μs', '~145 μs'),
            ('Standard', 'RFC 7748', 'FIPS 203', 'IETF Draft'),
            ('Maturity', '10+ years', 'Newly standardized', 'Emerging'),
        ],
        '667EEA',
        [1.5, 1.5, 1.5, 1.8])
    
    doc.add_paragraph()
    add_caption(doc, "Table 4.2: X25519 vs ML-KEM-768 vs Hybrid Comparison")
    
    doc.add_page_break()
    
    # =========================================================================
    # 5. FHE IMPLEMENTATION
    # =========================================================================
    doc.add_heading('5. Fully Homomorphic Encryption Implementation', level=1)
    
    p = doc.add_paragraph()
    p.add_run(
        "Fully Homomorphic Encryption (FHE) enables computation on encrypted data without decryption, "
        "allowing privacy-preserving analytics on sensitive information. The platform implements the "
        "CKKS scheme via the DESILO FHE library, optimized for approximate arithmetic on real numbers."
    )
    
    doc.add_heading('5.1 CKKS Scheme Overview', level=2)
    
    p = doc.add_paragraph()
    p.add_run(
        "The CKKS (Cheon-Kim-Kim-Song) scheme, published in ASIACRYPT 2017, supports approximate "
        "arithmetic operations on encrypted complex numbers. Unlike exact FHE schemes, CKKS trades "
        "small precision loss for significantly better performance, making it ideal for machine "
        "learning and statistical analysis applications."
    )
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Key advantages of CKKS include:")
    
    ckks_advantages = [
        "Native support for floating-point operations (addition, multiplication)",
        "Efficient SIMD-style parallelism via slot packing (8,192 parallel operations)",
        "Rescaling operation for noise management after multiplications",
        "Optional bootstrapping for unlimited computation depth",
        "GPU acceleration support for 5-6× performance improvement",
    ]
    for adv in ckks_advantages:
        doc.add_paragraph(f"• {adv}", style='List Bullet')
    
    doc.add_heading('5.2 DESILO FHE Configuration', level=2)
    
    p = doc.add_paragraph()
    p.add_run(
        "The platform uses optimized CKKS parameters balancing security, precision, and performance. "
        "These parameters provide 128-bit security equivalent while supporting 4 multiplicative depths."
    )
    
    doc.add_paragraph()
    
    add_styled_table(doc,
        ['Parameter', 'Value', 'Description', 'Impact'],
        [
            ('poly_degree', '16,384', 'Polynomial ring dimension (N)', 'Security vs performance'),
            ('coeff_mod_bit_sizes', '[60,40,40,40,60]', 'Coefficient modulus chain', 'Computation depth'),
            ('scale', '2^40', 'Encoding scale factor', 'Precision vs range'),
            ('max_mult_depth', '4', 'Maximum multiplicative depth', 'Circuit complexity'),
            ('slot_count', '8,192', 'Number of plaintext slots', 'Parallelism'),
            ('security_level', '128-bit', 'Equivalent symmetric security', 'Protection level'),
        ],
        '805AD5',
        [1.5, 1.5, 1.7, 1.5])
    
    doc.add_paragraph()
    add_caption(doc, "Table 5.1: CKKS Parameter Configuration")
    
    doc.add_page_break()
    
    doc.add_heading('5.3 Supported Operations', level=2)
    
    p = doc.add_paragraph()
    p.add_run(
        "The FHE engine supports a comprehensive set of homomorphic operations. Each operation has "
        "an associated depth cost that affects the remaining computation budget."
    )
    
    doc.add_paragraph()
    
    add_styled_table(doc,
        ['Operation', 'Input Types', 'Output', 'Depth Cost', 'Notes'],
        [
            ('Encrypt', 'Plaintext vector', 'Ciphertext', '0', 'Uses public key'),
            ('Decrypt', 'Ciphertext', 'Plaintext vector', '0', 'Uses secret key'),
            ('Add (CT+CT)', 'Two ciphertexts', 'Ciphertext', '0', 'No depth increase'),
            ('Add (CT+PT)', 'CT + plaintext', 'Ciphertext', '0', 'Efficient operation'),
            ('Multiply (scalar)', 'CT × number', 'Ciphertext', '0', 'No relinearization'),
            ('Multiply (CT×CT)', 'Two ciphertexts', 'Ciphertext', '1', 'Requires relinearization'),
            ('Rotate', 'Ciphertext, steps', 'Ciphertext', '0', 'Uses rotation keys'),
            ('Bootstrap', 'Ciphertext', 'Ciphertext', 'Reset', 'Refreshes noise budget'),
        ],
        '9F7AEA',
        [1.3, 1.3, 1.2, 0.8, 1.7])
    
    doc.add_paragraph()
    add_caption(doc, "Table 5.2: Supported FHE Operations")
    
    doc.add_heading('5.4 Use Case Example: Healthcare Analytics', level=2)
    
    p = doc.add_paragraph()
    p.add_run(
        "A typical healthcare analytics workflow using FHE:"
    )
    
    workflow_steps = [
        "Hospital encrypts patient vital signs (blood pressure, heart rate) using FHE public key",
        "Encrypted data is sent to cloud analytics provider",
        "Provider computes statistics (mean, variance, trends) on encrypted data",
        "Encrypted results are returned to hospital",
        "Hospital decrypts results using private key - provider never sees raw data",
    ]
    
    for i, step in enumerate(workflow_steps, 1):
        doc.add_paragraph(f"{i}. {step}")
    
    doc.add_page_break()
    
    # =========================================================================
    # 6. KUBERNETES DEPLOYMENT
    # =========================================================================
    doc.add_heading('6. Kubernetes Deployment', level=1)
    
    p = doc.add_paragraph()
    p.add_run(
        "The platform includes a production-ready Helm chart for Kubernetes deployment, supporting "
        "horizontal pod autoscaling, GPU workers, distributed caching, and comprehensive monitoring. "
        "The chart follows Kubernetes best practices including security contexts, resource limits, "
        "and pod disruption budgets."
    )
    
    doc.add_heading('6.1 Architecture Overview', level=2)
    
    # K8s architecture diagram as table
    k8s_arch = [
        ("INGRESS (TLS Termination)", "4299E1"),
        ("SERVICE (ClusterIP:8000)", "48BB78"),
        ("API DEPLOYMENT (2-10 Pods) ←→ HPA", "E2E8F0"),
        ("GPU WORKER (Optional, NVIDIA)", "805AD5"),
        ("REDIS | PROMETHEUS | GRAFANA", "E53E3E"),
    ]
    
    k8s_table = doc.add_table(rows=len(k8s_arch), cols=1)
    k8s_table.style = 'Table Grid'
    for i, (text, color) in enumerate(k8s_arch):
        cell = k8s_table.rows[i].cells[0]
        cell.text = text
        set_cell_shading(cell, color)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                if color not in ['E2E8F0']:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    
    doc.add_paragraph()
    add_caption(doc, "Figure 6.1: Kubernetes Deployment Architecture")
    
    doc.add_heading('6.2 Helm Chart Features', level=2)
    
    features = [
        ("Horizontal Pod Autoscaling (HPA)", "Automatically scales API replicas between 2 and 10 "
         "based on CPU utilization (70%) and memory utilization (80%) thresholds."),
        ("GPU Worker Support", "Optional deployment of GPU-accelerated workers using NVIDIA device "
         "plugin, with tolerations for GPU-specific node taints."),
        ("Redis Integration", "Bitnami Redis chart as dependency for distributed caching, supporting "
         "both standalone and replication architectures."),
        ("Prometheus Integration", "ServiceMonitor custom resource for automatic service discovery, "
         "with pre-configured scrape intervals and relabeling rules."),
        ("Network Policies", "Ingress/egress rules limiting traffic to authorized namespaces and CIDR "
         "blocks, implementing zero-trust networking principles."),
        ("Pod Disruption Budget", "Ensures minimum availability during rolling updates and node "
         "maintenance operations."),
        ("Ingress Configuration", "NGINX ingress with TLS termination, SSL redirect, and "
         "cert-manager integration for automatic certificate management."),
    ]
    for title, desc in features:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f"{title}: ")
        run.font.bold = True
        p.add_run(desc)
    
    doc.add_page_break()
    
    doc.add_heading('6.3 Configuration Reference', level=2)
    
    add_styled_table(doc,
        ['Parameter', 'Default', 'Description'],
        [
            ('api.replicaCount', '2', 'Initial API pod replicas'),
            ('api.image.repository', 'pqc-fhe-api', 'Container image repository'),
            ('api.resources.limits.cpu', '2000m', 'CPU limit per pod'),
            ('api.resources.limits.memory', '4Gi', 'Memory limit per pod'),
            ('api.resources.requests.cpu', '500m', 'CPU request per pod'),
            ('api.resources.requests.memory', '1Gi', 'Memory request per pod'),
            ('api.autoscaling.enabled', 'true', 'Enable HPA'),
            ('api.autoscaling.minReplicas', '2', 'Minimum replicas'),
            ('api.autoscaling.maxReplicas', '10', 'Maximum replicas'),
            ('api.autoscaling.targetCPU', '70', 'Target CPU utilization (%)'),
            ('gpuWorker.enabled', 'false', 'Enable GPU workers'),
            ('gpuWorker.resources.nvidia.com/gpu', '1', 'GPUs per worker'),
            ('redis.enabled', 'true', 'Enable Redis cache'),
            ('redis.master.persistence.size', '8Gi', 'Redis storage size'),
            ('prometheus.enabled', 'true', 'Enable Prometheus'),
            ('prometheus.server.retention', '15d', 'Metrics retention period'),
            ('networkPolicy.enabled', 'true', 'Enable network policies'),
            ('podDisruptionBudget.enabled', 'true', 'Enable PDB'),
            ('podDisruptionBudget.minAvailable', '1', 'Minimum available pods'),
        ],
        '4299E1',
        [2.5, 1.0, 3.0])
    
    doc.add_paragraph()
    add_caption(doc, "Table 6.1: Helm Chart Configuration Parameters")
    
    doc.add_heading('6.4 Quick Start Commands', level=2)
    
    commands = [
        ("Install dependencies", "helm dependency update ./kubernetes/helm/pqc-fhe"),
        ("Basic install", "helm install pqc-fhe ./kubernetes/helm/pqc-fhe --namespace pqc-fhe --create-namespace"),
        ("With GPU workers", "helm install pqc-fhe ./kubernetes/helm/pqc-fhe --set gpuWorker.enabled=true"),
        ("Upgrade", "helm upgrade pqc-fhe ./kubernetes/helm/pqc-fhe"),
        ("Uninstall", "helm uninstall pqc-fhe"),
    ]
    
    for desc, cmd in commands:
        p = doc.add_paragraph()
        run = p.add_run(f"{desc}: ")
        run.font.bold = True
        run = p.add_run(cmd)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    
    doc.add_page_break()
    
    # =========================================================================
    # 7. MONITORING
    # =========================================================================
    doc.add_heading('7. Monitoring and Observability', level=1)
    
    p = doc.add_paragraph()
    p.add_run(
        "The platform integrates comprehensive monitoring capabilities using the Prometheus ecosystem. "
        "Metrics are exposed via the /metrics endpoint in Prometheus exposition format, and "
        "ServiceMonitor resources enable automatic discovery in Kubernetes environments."
    )
    
    doc.add_heading('7.1 Metrics Architecture', level=2)
    
    p = doc.add_paragraph()
    p.add_run(
        "The monitoring stack consists of three main components: Prometheus for metrics collection "
        "and alerting, Grafana for visualization, and AlertManager for notification routing. All "
        "components are pre-configured in the Helm chart."
    )
    
    doc.add_heading('7.2 Exposed Metrics', level=2)
    
    add_styled_table(doc,
        ['Metric Name', 'Type', 'Labels', 'Description'],
        [
            ('http_requests_total', 'Counter', 'method, endpoint, status', 'Total HTTP requests'),
            ('http_request_duration_seconds', 'Histogram', 'method, endpoint', 'Request latency distribution'),
            ('http_request_size_bytes', 'Histogram', 'method, endpoint', 'Request body size'),
            ('http_response_size_bytes', 'Histogram', 'method, endpoint', 'Response body size'),
            ('pqc_keygen_duration_seconds', 'Histogram', 'algorithm', 'Key generation time'),
            ('pqc_encapsulate_duration_seconds', 'Histogram', 'algorithm', 'Encapsulation time'),
            ('pqc_sign_duration_seconds', 'Histogram', 'algorithm', 'Signing time'),
            ('fhe_encrypt_duration_seconds', 'Histogram', 'slot_count', 'FHE encryption time'),
            ('fhe_decrypt_duration_seconds', 'Histogram', 'slot_count', 'FHE decryption time'),
            ('fhe_operation_duration_seconds', 'Histogram', 'operation', 'FHE operation time'),
            ('ciphertext_store_size', 'Gauge', '-', 'Stored ciphertexts count'),
            ('keypair_store_size', 'Gauge', 'type', 'Stored keypairs count'),
        ],
        'ED8936',
        [2.0, 0.8, 1.5, 2.0])
    
    doc.add_paragraph()
    add_caption(doc, "Table 7.1: Prometheus Metrics Reference")
    
    doc.add_page_break()
    
    doc.add_heading('7.3 Pre-configured Alerts', level=2)
    
    p = doc.add_paragraph()
    p.add_run(
        "The platform includes pre-configured Prometheus alerting rules for common failure scenarios. "
        "These alerts are defined in the servicemonitor.yaml template and can be customized via Helm values."
    )
    
    doc.add_paragraph()
    
    add_styled_table(doc,
        ['Alert Name', 'Condition', 'Duration', 'Severity', 'Action'],
        [
            ('PQCFHEHighErrorRate', 'Error rate > 5%', '5 min', 'Critical', 'Page on-call'),
            ('PQCFHEHighLatency', 'p95 latency > 5s', '5 min', 'Warning', 'Investigate'),
            ('PQCFHEPodNotReady', 'Replicas < desired', '10 min', 'Warning', 'Check pods'),
            ('PQCFHESlowEncryption', 'p95 encrypt > 10s', '5 min', 'Warning', 'Scale GPU'),
            ('PQCFHEGPUMemoryHigh', 'GPU memory > 90%', '5 min', 'Warning', 'Add capacity'),
            ('PQCFHEGPUUnderutilized', 'GPU util < 10%', '1 hour', 'Info', 'Reduce GPUs'),
        ],
        'E53E3E',
        [1.5, 1.2, 0.8, 0.8, 1.2])
    
    doc.add_paragraph()
    add_caption(doc, "Table 7.2: Pre-configured Prometheus Alerts")
    
    doc.add_heading('7.4 Grafana Dashboards', level=2)
    
    p = doc.add_paragraph()
    p.add_run(
        "Recommended Grafana dashboards for monitoring the platform:"
    )
    
    dashboards = [
        "API Overview: Request rates, latencies, error rates by endpoint",
        "Cryptographic Operations: PQC key generation, encapsulation, signing times",
        "FHE Performance: Encryption/decryption times, bootstrap frequency",
        "Resource Utilization: CPU, memory, GPU usage per pod",
        "Kubernetes: Pod health, HPA status, network I/O",
    ]
    
    for dashboard in dashboards:
        doc.add_paragraph(f"• {dashboard}", style='List Bullet')
    
    doc.add_page_break()
    
    # =========================================================================
    # 8. LOGGING
    # =========================================================================
    doc.add_heading('8. Logging System', level=1)
    
    p = doc.add_paragraph()
    p.add_run(
        "The platform implements enterprise-grade file-based logging with automatic rotation, "
        "separate log streams for different purposes, and configurable verbosity levels. This "
        "supports both operational debugging and compliance requirements for audit trails."
    )
    
    doc.add_heading('8.1 Logging Architecture', level=2)
    
    # Logging diagram
    log_arch = [
        ("API SERVER (FastAPI Logger)", "3182CE"),
        ("↓", "FFFFFF"),
        ("pqc_fhe_server.log (All logs, 10MB × 5)", "38A169"),
        ("pqc_fhe_error.log (Errors only, 10MB × 3)", "E53E3E"),
        ("pqc_fhe_access.log (HTTP access, 10MB × 5)", "805AD5"),
        ("+ Console Output (stdout)", "2D3748"),
    ]
    
    log_table = doc.add_table(rows=len(log_arch), cols=1)
    log_table.style = 'Table Grid'
    for i, (text, color) in enumerate(log_arch):
        cell = log_table.rows[i].cells[0]
        cell.text = text
        if color != 'FFFFFF':
            set_cell_shading(cell, color)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(9)
                if color not in ['FFFFFF', 'E2E8F0']:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.bold = True
    
    doc.add_paragraph()
    add_caption(doc, "Figure 8.1: File-Based Logging Architecture")
    
    doc.add_heading('8.2 Log Files', level=2)
    
    add_styled_table(doc,
        ['File Name', 'Max Size', 'Backups', 'Level', 'Content Description'],
        [
            ('pqc_fhe_server.log', '10 MB', '5', 'INFO+', 'All server operations and events'),
            ('pqc_fhe_error.log', '10 MB', '3', 'ERROR+', 'Errors and exceptions only'),
            ('pqc_fhe_access.log', '10 MB', '5', 'INFO', 'HTTP request/response logs'),
        ],
        '667EEA',
        [1.8, 0.7, 0.7, 0.7, 2.4])
    
    doc.add_paragraph()
    add_caption(doc, "Table 8.1: Log File Configuration")
    
    doc.add_heading('8.3 Log Format', level=2)
    
    p = doc.add_paragraph()
    run = p.add_run("File log format ")
    p.add_run("(includes source location for debugging):")
    
    doc.add_paragraph()
    code_p = doc.add_paragraph()
    run = code_p.add_run("2025-12-30 12:00:00 - api.server - INFO - [server.py:123] - Request processed")
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Console log format ")
    p.add_run("(compact for readability):")
    
    doc.add_paragraph()
    code_p = doc.add_paragraph()
    run = code_p.add_run("2025-12-30 12:00:00 - api.server - INFO - Request processed")
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    
    doc.add_heading('8.4 Configuration', level=2)
    
    p = doc.add_paragraph()
    p.add_run(
        "Log verbosity can be configured via the LOG_LEVEL environment variable. Supported levels "
        "in order of increasing verbosity are: CRITICAL, ERROR, WARNING, INFO (default), DEBUG."
    )
    
    doc.add_paragraph()
    code_p = doc.add_paragraph()
    run = code_p.add_run("export LOG_LEVEL=DEBUG")
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    
    doc.add_page_break()
    
    # =========================================================================
    # 9. API REFERENCE
    # =========================================================================
    doc.add_heading('9. API Reference', level=1)
    
    p = doc.add_paragraph()
    p.add_run(
        "The platform exposes a comprehensive REST API with automatic OpenAPI documentation "
        "at /docs (Swagger UI) and /redoc (ReDoc). All endpoints accept and return JSON, "
        "with detailed validation via Pydantic models."
    )
    
    doc.add_heading('9.1 Hybrid Key Exchange Endpoints', level=2)
    
    add_styled_table(doc,
        ['Endpoint', 'Method', 'Description', 'Request Body'],
        [
            ('/pqc/hybrid/keypair', 'POST', 'Generate hybrid keypair', '{"kem_algorithm": "ML-KEM-768"}'),
            ('/pqc/hybrid/encapsulate', 'POST', 'Hybrid encapsulation', '{"keypair_id": "..."}'),
            ('/pqc/hybrid/decapsulate', 'POST', 'Hybrid decapsulation', '{"keypair_id", "ephemeral_public", "ciphertext"}'),
            ('/pqc/hybrid/compare', 'GET', 'Algorithm comparison', '-'),
            ('/pqc/hybrid/migration-strategy', 'GET', 'Migration roadmap', '-'),
            ('/pqc/hybrid/keypairs', 'GET', 'List stored keypairs', '-'),
        ],
        '667EEA',
        [1.8, 0.6, 1.5, 2.4])
    
    doc.add_paragraph()
    add_caption(doc, "Table 9.1: Hybrid Key Exchange API Endpoints")
    
    doc.add_heading('9.2 PQC Endpoints', level=2)
    
    add_styled_table(doc,
        ['Endpoint', 'Method', 'Description', 'Request Body'],
        [
            ('/pqc/algorithms', 'GET', 'List available algorithms', '-'),
            ('/pqc/kem/keypair', 'POST', 'Generate ML-KEM keypair', '{"algorithm": "ML-KEM-768"}'),
            ('/pqc/kem/encapsulate', 'POST', 'Encapsulate shared secret', '{"keypair_id": "..."}'),
            ('/pqc/kem/decapsulate', 'POST', 'Decapsulate shared secret', '{"keypair_id", "ciphertext"}'),
            ('/pqc/sig/keypair', 'POST', 'Generate ML-DSA keypair', '{"algorithm": "ML-DSA-65"}'),
            ('/pqc/sig/sign', 'POST', 'Sign message', '{"keypair_id", "message"}'),
            ('/pqc/sig/verify', 'POST', 'Verify signature', '{"public_key", "message", "signature"}'),
        ],
        '48BB78',
        [1.8, 0.6, 1.6, 2.3])
    
    doc.add_paragraph()
    add_caption(doc, "Table 9.2: PQC API Endpoints")
    
    doc.add_heading('9.3 FHE Endpoints', level=2)
    
    add_styled_table(doc,
        ['Endpoint', 'Method', 'Description', 'Request Body'],
        [
            ('/fhe/encrypt', 'POST', 'Encrypt numeric vector', '{"data": [1.0, 2.0, ...]}'),
            ('/fhe/decrypt', 'POST', 'Decrypt ciphertext', '{"ciphertext_id": "..."}'),
            ('/fhe/add', 'POST', 'Homomorphic addition', '{"ct_id_a", "ct_id_b"}'),
            ('/fhe/multiply', 'POST', 'Homomorphic multiplication', '{"ct_id", "scalar"}'),
            ('/fhe/ciphertexts', 'GET', 'List stored ciphertexts', '-'),
        ],
        '805AD5',
        [1.8, 0.6, 1.6, 2.3])
    
    doc.add_paragraph()
    add_caption(doc, "Table 9.3: FHE API Endpoints")
    
    doc.add_page_break()
    
    doc.add_heading('9.4 Enterprise Demo Endpoints', level=2)
    
    add_styled_table(doc,
        ['Endpoint', 'Method', 'Description', 'Data Source'],
        [
            ('/enterprise/healthcare', 'GET', 'Healthcare vital signs demo', 'VitalDB'),
            ('/enterprise/finance', 'GET', 'Financial portfolio demo', 'Yahoo Finance'),
            ('/enterprise/iot', 'GET', 'IoT smart grid demo', 'Synthetic'),
            ('/enterprise/blockchain', 'GET', 'Blockchain transaction demo', 'Ethereum RPC'),
        ],
        'DD6B20',
        [2.0, 0.6, 1.8, 1.9])
    
    doc.add_paragraph()
    add_caption(doc, "Table 9.4: Enterprise Demo Endpoints")
    
    doc.add_heading('9.5 System Endpoints', level=2)
    
    add_styled_table(doc,
        ['Endpoint', 'Method', 'Description'],
        [
            ('/health', 'GET', 'Health check endpoint'),
            ('/metrics', 'GET', 'Prometheus metrics'),
            ('/docs', 'GET', 'Swagger UI documentation'),
            ('/redoc', 'GET', 'ReDoc documentation'),
            ('/ui', 'GET', 'Web UI interface'),
        ],
        '718096',
        [2.0, 0.8, 3.5])
    
    doc.add_page_break()
    
    # =========================================================================
    # 10. ENTERPRISE USE CASES
    # =========================================================================
    doc.add_heading('10. Enterprise Use Cases', level=1)
    
    p = doc.add_paragraph()
    p.add_run(
        "The platform supports multiple enterprise use cases across regulated industries, "
        "demonstrating practical applications of quantum-resistant cryptography and "
        "privacy-preserving computation."
    )
    
    doc.add_heading('10.1 Healthcare: HIPAA-Compliant Analytics', level=2)
    
    p = doc.add_paragraph()
    p.add_run(
        "Healthcare organizations can analyze patient vital signs without exposing Protected Health "
        "Information (PHI). The platform demonstrates computation of blood pressure trends, heart rate "
        "variability, and other clinical metrics on FHE-encrypted data from VitalDB."
    )
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Workflow: ")
    run.font.bold = True
    p.add_run(
        "Hospital encrypts PHI → Cloud provider computes statistics on encrypted data → "
        "Hospital decrypts results. Provider never accesses raw patient data."
    )
    
    doc.add_heading('10.2 Finance: Confidential Portfolio Analysis', level=2)
    
    p = doc.add_paragraph()
    p.add_run(
        "Investment firms can perform growth projections on encrypted portfolio values using live "
        "market data from Yahoo Finance. Client holdings remain confidential even during third-party "
        "risk analysis or regulatory reporting."
    )
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Workflow: ")
    run.font.bold = True
    p.add_run(
        "Firm encrypts portfolio → Analytics provider computes VaR, projections → "
        "Firm decrypts results. Provider cannot see actual positions or values."
    )
    
    doc.add_heading('10.3 IoT: Secure Smart Grid Analytics', level=2)
    
    p = doc.add_paragraph()
    p.add_run(
        "Utility companies can aggregate encrypted smart meter readings for demand forecasting "
        "without accessing individual household consumption patterns. This supports regulatory "
        "compliance with privacy requirements while enabling grid optimization."
    )
    
    doc.add_heading('10.4 Blockchain: Quantum-Resistant Transactions', level=2)
    
    p = doc.add_paragraph()
    p.add_run(
        "Cryptocurrency platforms can migrate from ECDSA to ML-DSA signatures, protecting transaction "
        "integrity against future quantum attacks. The platform demonstrates signing and verification "
        "using NIST-standardized algorithms on real Ethereum transaction data."
    )
    
    doc.add_page_break()
    
    # =========================================================================
    # 11. SECURITY ANALYSIS
    # =========================================================================
    doc.add_heading('11. Security Analysis', level=1)
    
    p = doc.add_paragraph()
    p.add_run(
        "The platform implements multiple layers of security based on NIST guidelines and "
        "industry best practices. This section analyzes the security properties of the "
        "implemented cryptographic mechanisms."
    )
    
    doc.add_heading('11.1 NIST Security Levels', level=2)
    
    p = doc.add_paragraph()
    p.add_run(
        "NIST defines five security levels based on the computational effort required to break "
        "the cryptographic scheme. The platform supports Levels 1, 3, and 5."
    )
    
    doc.add_paragraph()
    
    add_styled_table(doc,
        ['Level', 'Classical Equivalent', 'Algorithms Supported'],
        [
            ('Level 1 (128-bit)', 'AES-128', 'ML-KEM-512'),
            ('Level 3 (192-bit)', 'AES-192', 'ML-KEM-768, ML-DSA-65, Hybrid'),
            ('Level 5 (256-bit)', 'AES-256', 'ML-KEM-1024, ML-DSA-87'),
        ],
        'C53030',
        [1.5, 1.5, 3.0])
    
    doc.add_paragraph()
    add_caption(doc, "Figure 11.1: NIST Security Levels and Algorithm Mapping")
    
    doc.add_heading('11.2 Threat Model', level=2)
    
    add_styled_table(doc,
        ['Threat', 'Mitigation', 'Algorithm'],
        [
            ('Quantum key recovery', 'Lattice-based hardness (MLWE)', 'ML-KEM'),
            ('Quantum signature forgery', 'Module-LWE security', 'ML-DSA'),
            ('Harvest now, decrypt later', 'Hybrid key exchange', 'X25519 + ML-KEM'),
            ('Side-channel attacks', 'Constant-time implementations', 'All'),
            ('Implementation bugs', 'Defense in depth (hybrid)', 'Combined'),
            ('Data exposure in transit', 'PQC-secured TLS', 'Hybrid TLS'),
            ('Data exposure at rest', 'FHE computation', 'CKKS'),
        ],
        'C53030',
        [2.0, 2.2, 2.0])
    
    doc.add_paragraph()
    add_caption(doc, "Table 11.1: Security Threat Model")
    
    doc.add_page_break()
    
    # =========================================================================
    # 12. PERFORMANCE BENCHMARKS
    # =========================================================================
    doc.add_heading('12. Performance Benchmarks', level=1)
    
    p = doc.add_paragraph()
    p.add_run(
        "Performance measurements were conducted on an Intel Core i7-12700H processor with "
        "32GB RAM and NVIDIA RTX 4090 GPU. Results represent average values over 1000 iterations."
    )
    
    doc.add_heading('12.1 Hybrid Key Exchange Performance', level=2)
    
    add_styled_table(doc,
        ['Operation', 'X25519', 'ML-KEM-768', 'Hybrid', 'Overhead'],
        [
            ('Key Generation', '18 μs', '25 μs', '43 μs', '+0 μs'),
            ('Encapsulation', '20 μs', '30 μs', '52 μs', '+2 μs'),
            ('Decapsulation', '20 μs', '28 μs', '50 μs', '+2 μs'),
            ('Total Round-Trip', '58 μs', '83 μs', '145 μs', '+4 μs'),
        ],
        '3182CE',
        [1.4, 1.0, 1.0, 0.9, 0.9])
    
    doc.add_paragraph()
    add_caption(doc, "Table 12.1: Hybrid Key Exchange Performance")
    
    doc.add_heading('12.2 FHE Operations Performance', level=2)
    
    add_styled_table(doc,
        ['Operation', 'CPU Time', 'GPU Time (RTX 4090)', 'Speedup'],
        [
            ('Key Generation', '2.5 s', '0.8 s', '3.1×'),
            ('Encrypt (8192 slots)', '15 ms', '3 ms', '5.0×'),
            ('Decrypt', '10 ms', '2 ms', '5.0×'),
            ('Add (CT + CT)', '0.5 ms', '0.1 ms', '5.0×'),
            ('Multiply (CT × scalar)', '2 ms', '0.3 ms', '6.7×'),
            ('Multiply (CT × CT)', '50 ms', '8 ms', '6.3×'),
            ('Bootstrap', '15 s', '2.5 s', '6.0×'),
        ],
        '805AD5',
        [1.6, 1.3, 1.5, 0.8])
    
    doc.add_paragraph()
    add_caption(doc, "Table 12.2: FHE Operations Performance (CPU vs GPU)")
    
    doc.add_heading('12.3 Throughput', level=2)
    
    p = doc.add_paragraph()
    p.add_run(
        "Under load testing with 100 concurrent users, the platform achieved:"
    )
    
    throughput = [
        "Hybrid key exchange: 6,500 operations/second",
        "PQC signatures: 5,000 sign/verify operations/second",
        "FHE encryption: 60 operations/second (CPU), 300 operations/second (GPU)",
        "API response time (p95): 50ms for PQC, 200ms for FHE",
    ]
    
    for t in throughput:
        doc.add_paragraph(f"• {t}", style='List Bullet')
    
    doc.add_page_break()
    
    # =========================================================================
    # 13. FUTURE ROADMAP
    # =========================================================================
    doc.add_heading('13. Future Roadmap', level=1)
    
    p = doc.add_paragraph()
    p.add_run(
        "The platform development roadmap extends through 2026, with planned features including "
        "additional cryptographic algorithms, FIPS validation, and hardware security module integration."
    )
    
    doc.add_paragraph()
    
    add_styled_table(doc,
        ['Version', 'Timeline', 'Major Features'],
        [
            ('v2.4.0', 'Q1 2025', 'SLH-DSA (FIPS 205) hash-based signatures'),
            ('v2.5.0', 'Q2 2025', 'Native TLS 1.3 hybrid integration'),
            ('v2.6.0', 'Q3 2025', 'Multi-party computation (MPC) framework'),
            ('v3.0.0', 'Q4 2025', 'FIPS validation and CMVP certification'),
            ('v3.1.0', 'Q1 2026', 'Hardware security module (HSM) integration'),
            ('v3.2.0', 'Q2 2026', 'Zero-knowledge proof support'),
        ],
        '38A169',
        [1.0, 1.0, 4.5])
    
    doc.add_paragraph()
    add_caption(doc, "Table 13.1: Development Roadmap")
    
    doc.add_page_break()
    
    # =========================================================================
    # APPENDIX A
    # =========================================================================
    doc.add_heading('Appendix A: Helm Chart Configuration', level=1)
    
    p = doc.add_paragraph()
    p.add_run("Example values.yaml configuration for production deployment:")
    
    doc.add_paragraph()
    
    helm_config = """api:
  replicaCount: 3
  resources:
    limits:
      cpu: 4000m
      memory: 8Gi
    requests:
      cpu: 1000m
      memory: 2Gi
  autoscaling:
    enabled: true
    minReplicas: 3
    maxReplicas: 20
    targetCPUUtilizationPercentage: 60

gpuWorker:
  enabled: true
  replicaCount: 2
  resources:
    limits:
      nvidia.com/gpu: 1
      memory: 32Gi

crypto:
  pqc:
    kemAlgorithm: ML-KEM-768
    signatureAlgorithm: ML-DSA-65
  fhe:
    mode: gpu
    useBootstrap: true

redis:
  enabled: true
  architecture: replication
  master:
    persistence:
      size: 16Gi

prometheus:
  enabled: true
  server:
    retention: 30d
    persistentVolume:
      size: 100Gi"""
    
    for line in helm_config.split('\n'):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(8)
    
    doc.add_page_break()
    
    # =========================================================================
    # APPENDIX B
    # =========================================================================
    doc.add_heading('Appendix B: Prometheus Alerting Rules', level=1)
    
    p = doc.add_paragraph()
    p.add_run("Example PrometheusRule configuration:")
    
    doc.add_paragraph()
    
    prom_rules = """groups:
  - name: pqc-fhe-api
    rules:
      - alert: PQCFHEHighErrorRate
        expr: |
          (sum(rate(http_requests_total{status=~"5.."}[5m]))
           / sum(rate(http_requests_total[5m]))) > 0.05
        for: 5m
        labels:
          severity: critical
        annotations:
          summary: "High error rate in PQC-FHE API"
          
      - alert: PQCFHEHighLatency
        expr: |
          histogram_quantile(0.95,
            sum(rate(http_request_duration_seconds_bucket[5m])) by (le)
          ) > 5
        for: 5m
        labels:
          severity: warning
        annotations:
          summary: "High latency in PQC-FHE API"
          
      - alert: PQCFHESlowEncryption
        expr: |
          histogram_quantile(0.95,
            sum(rate(fhe_encryption_duration_seconds_bucket[5m])) by (le)
          ) > 10
        for: 5m
        labels:
          severity: warning
        annotations:
          summary: "FHE encryption operations are slow" """
    
    for line in prom_rules.split('\n'):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(8)
    
    doc.add_page_break()
    
    # =========================================================================
    # REFERENCES
    # =========================================================================
    doc.add_heading('References', level=1)
    
    references = [
        "[1] NIST. FIPS 203: Module-Lattice-Based Key-Encapsulation Mechanism Standard. "
        "National Institute of Standards and Technology, August 2024. "
        "https://csrc.nist.gov/pubs/fips/203/final",
        
        "[2] NIST. FIPS 204: Module-Lattice-Based Digital Signature Standard. "
        "National Institute of Standards and Technology, August 2024. "
        "https://csrc.nist.gov/pubs/fips/204/final",
        
        "[3] NIST. FIPS 205: Stateless Hash-Based Digital Signature Standard. "
        "National Institute of Standards and Technology, August 2024. "
        "https://csrc.nist.gov/pubs/fips/205/final",
        
        "[4] NIST. IR 8547: Transition to Post-Quantum Cryptography Standards. "
        "https://csrc.nist.gov/pubs/ir/8547/final",
        
        "[5] IETF. draft-ietf-tls-ecdhe-mlkem: Hybrid Key Exchange for TLS 1.3. "
        "https://datatracker.ietf.org/doc/draft-ietf-tls-ecdhe-mlkem/",
        
        "[6] Bernstein DJ, Lange T. RFC 7748: Elliptic Curves for Security. "
        "https://datatracker.ietf.org/doc/html/rfc7748",
        
        "[7] Cheon JH, Kim A, Kim M, Song Y. Homomorphic Encryption for Arithmetic "
        "of Approximate Numbers. ASIACRYPT 2017. DOI: 10.1007/978-3-319-70694-8_15",
        
        "[8] DESILO. DESILO FHE Library Documentation. https://fhe.desilo.dev/latest/",
        
        "[9] Open Quantum Safe. liboqs-python: Python 3 wrapper for liboqs. "
        "https://github.com/open-quantum-safe/liboqs-python",
        
        "[10] Lee HC, et al. VitalDB Database. Scientific Data 9, 279 (2022). "
        "DOI: 10.1038/s41597-022-01411-5",
        
        "[11] Kubernetes. Helm Documentation. https://helm.sh/docs/",
        
        "[12] Prometheus. Prometheus Operator. https://prometheus-operator.dev/",
    ]
    
    for ref in references:
        p = doc.add_paragraph()
        p.add_run(ref)
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.3)
        p.paragraph_format.space_after = Pt(6)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Footer
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 50)
    run.font.color.rgb = RGBColor(0xA0, 0xAE, 0xC0)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Document Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PQC-FHE Integration Platform v2.3.5 Enterprise Edition")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("© 2025 - MIT License")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
    
    # Save
    doc.save("PQC_FHE_Technical_Report_v2.3.5_Enterprise_Full.docx")
    print("Word document generated: PQC_FHE_Technical_Report_v2.3.5_Enterprise_Full.docx")


if __name__ == "__main__":
    create_report()
